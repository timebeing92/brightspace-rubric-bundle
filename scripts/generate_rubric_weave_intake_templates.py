#!/usr/bin/env python3
"""Generate the versioned Rubric Weave intake templates deterministically.

The templates are presentation assets over the accepted Workbench producer at
``7c5140545548c89a254ac4502cfdd7ee6fb44255``.  This generator does not parse,
normalize, score, or package rubrics; those semantics remain in
``rubric_authoring.py`` and ``make_rubric_package.py``.
"""

from __future__ import annotations

import argparse
import hashlib
import io
import json
import math
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT_DIR = (
    REPO_ROOT / "workspace/reference/templates/rubric-weave/v1"
)

TEMPLATE_VERSION = "v1"
TEMPLATE_SET = "rubric-weave-intake"
ACCEPTED_PRODUCER_COMMIT = "7c5140545548c89a254ac4502cfdd7ee6fb44255"
AUTHORING_CONTRACT = "coursecraft.rubric_authoring/1"
DOCX_NAME = "rubric-weave-intake-template.docx"
MARKDOWN_NAME = "rubric-weave-intake-template.md"
MANIFEST_NAME = "manifest.json"

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
MARKDOWN_MEDIA_TYPE = "text/markdown"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
MARGIN_DXA = 1440
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_DXA = 80
CELL_MARGIN_BOTTOM_DXA = 80
CELL_MARGIN_START_DXA = 120
CELL_MARGIN_END_DXA = 120
FONT_NAME = "Calibri"
HEADING_BLUE = "2E74B5"
HEADER_FILL = "E8EEF5"
TABLE_BORDER = "AAB7C4"
DOCX_DESIGN_PRESET = "compact_reference_guide"
# Named parser-form overrides: the sole title heading starts at 0 pt before,
# and the document has no decorative header/footer furniture.  This preserves
# the accepted title -> table -> ordinary instructions body sequence.
DOCX_TITLE_BEFORE_OVERRIDE_PT = 0
FIXED_PACKAGE_TIMESTAMP = (1980, 1, 1, 0, 0, 0)
FIXED_DOCUMENT_TIME = datetime(2000, 1, 1, 0, 0, 0)


@dataclass(frozen=True)
class Level:
    """One editable performance-level column."""

    name: str
    score: float


@dataclass(frozen=True)
class Criterion:
    """One editable criterion row."""

    name: str
    weight: float
    descriptions: tuple[str, ...]


@dataclass(frozen=True)
class TemplateSpec:
    """Substrate-neutral content used by both canonical template renderers."""

    title: str
    levels: tuple[Level, ...]
    criteria: tuple[Criterion, ...]


DEFAULT_SPEC = TemplateSpec(
    title="SYNTHETIC PRACTICE RUBRIC - REPLACE BEFORE USE",
    levels=(
        Level("Ready to Share", 100),
        Level("Needs Revision", 60),
        Level("Not Yet Demonstrated", 0),
    ),
    criteria=(
        Criterion(
            "Evidence use",
            40,
            (
                "Synthetic response uses specific practice evidence and explains why it matters.",
                "Synthetic response uses some practice evidence, but the explanation is incomplete.",
                "Synthetic response does not yet connect practice evidence to its claim.",
            ),
        ),
        Criterion(
            "Reasoning",
            35,
            (
                "Synthetic response makes a clear claim and follows a complete line of reasoning.",
                "Synthetic response has a recognizable claim, but part of the reasoning needs revision.",
                "Synthetic response does not yet provide a clear or supported line of reasoning.",
            ),
        ),
        Criterion(
            "Audience and next step",
            25,
            (
                "Synthetic response fits its practice audience and names a useful next step.",
                "Synthetic response partly addresses its audience or gives a general next step.",
                "Synthetic response does not yet address its audience or identify a next step.",
            ),
        ),
    ),
)


INSTRUCTION_PARAGRAPHS = (
    (
        "Edit the synthetic example.",
        "Replace the title, criteria, level names, numeric scores, weights, and "
        "descriptions. You may add or remove criterion rows and performance-level "
        "columns; keep one Criterion column, one criterion per row, and at least "
        "two uniquely named level columns.",
    ),
    (
        "Keep scoring explicit.",
        "Put exactly one numeric score in every level header, such as Ready to "
        "Share (100). The Weight column is structurally optional, but this example "
        "includes explicit positive weights totaling 100 so it needs no fallback.",
    ),
    (
        "Avoid ambiguous Word structure.",
        "Keep one simple rectangular table. Do not merge cells, nest tables, use "
        "multiple header bands, add row-number or notes columns, detach scoring "
        "into another table, or rely on a decorative layout. Correct ambiguous "
        "structure or use Markdown or JSON.",
    ),
    (
        "Correct or approve a fallback visibly.",
        "If preflight reports missing or ambiguous scores or weights, correct the "
        "source and run it again. Approve even spacing or equal weights only when "
        "you intentionally choose that fallback; the producer records the "
        "approval and never invents scoring silently.",
    ),
    (
        "Respect the Brightspace boundary.",
        "Weave builds and validates a rubric-only import package; that is not a "
        "Brightspace import or proof of Brightspace acceptance. Importing the "
        "package does not attach the rubric to an assignment, discussion, quiz, "
        "or grade item. Activity attachment is a separate manual step.",
    ),
)


def _format_number(value: float) -> str:
    number = float(value)
    if number.is_integer():
        return str(int(number))
    return format(number, ".12g")


def _validate_spec(spec: TemplateSpec) -> None:
    if not spec.title.strip():
        raise ValueError("The template title must not be empty.")
    if len(spec.levels) < 2:
        raise ValueError("A rubric template requires at least two levels.")
    if not spec.criteria:
        raise ValueError("A rubric template requires at least one criterion.")

    normalized_level_names = [level.name.strip().casefold() for level in spec.levels]
    if any(not name for name in normalized_level_names):
        raise ValueError("Every performance level requires a name.")
    if len(set(normalized_level_names)) != len(normalized_level_names):
        raise ValueError("Performance-level names must be unique.")

    scores = [float(level.score) for level in spec.levels]
    if any(not math.isfinite(score) or score < 0 or score > 100 for score in scores):
        raise ValueError("Performance-level scores must be finite values from 0 to 100.")
    if len(set(scores)) != len(scores):
        raise ValueError("Performance-level scores must be unique.")

    normalized_criterion_names = [
        criterion.name.strip().casefold() for criterion in spec.criteria
    ]
    if any(not name for name in normalized_criterion_names):
        raise ValueError("Every criterion requires a name.")
    if len(set(normalized_criterion_names)) != len(normalized_criterion_names):
        raise ValueError("Criterion names must be unique.")

    weights = [float(criterion.weight) for criterion in spec.criteria]
    if any(not math.isfinite(weight) or weight <= 0 for weight in weights):
        raise ValueError("Criterion weights must be positive finite values.")
    if not math.isclose(sum(weights), 100.0, rel_tol=0.0, abs_tol=1e-9):
        raise ValueError("Criterion weights must total 100.")
    if any(
        len(criterion.descriptions) != len(spec.levels)
        or any(not value.strip() for value in criterion.descriptions)
        for criterion in spec.criteria
    ):
        raise ValueError("Every criterion requires one non-empty description per level.")


def _instruction_text(label: str, text: str) -> str:
    return f"{label} {text}"


def _markdown_cell(value: str) -> str:
    return " ".join(value.replace("\\", "\\\\").replace("|", "\\|").split())


def render_markdown(spec: TemplateSpec = DEFAULT_SPEC) -> bytes:
    """Render one parser-ready Markdown table plus equivalent instructions."""

    _validate_spec(spec)
    headers = [
        "Criterion",
        "Weight",
        *(
            f"{level.name} ({_format_number(level.score)})"
            for level in spec.levels
        ),
    ]
    lines = [
        f"## {spec.title}",
        "",
        "| " + " | ".join(_markdown_cell(value) for value in headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for criterion in spec.criteria:
        row = [
            criterion.name,
            _format_number(criterion.weight),
            *criterion.descriptions,
        ]
        lines.append("| " + " | ".join(_markdown_cell(value) for value in row) + " |")
    for label, text in INSTRUCTION_PARAGRAPHS:
        lines.extend(["", f"**{label}** {text}"])
    return ("\n".join(lines) + "\n").encode("utf-8")


def _replace_children(parent: object, tag: str, replacement: object) -> None:
    for child in list(parent.findall(qn(tag))):
        parent.remove(child)
    parent.append(replacement)


def _set_run_font(run: object, size_pt: float, *, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), FONT_NAME)


def _set_style_font(style: object, size_pt: float, color: str) -> None:
    style.font.name = FONT_NAME
    style.font.size = Pt(size_pt)
    style.font.color.rgb = RGBColor.from_string(color)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), FONT_NAME)


def _set_table_cell_margins(table: object) -> None:
    margins = OxmlElement("w:tblCellMar")
    for tag, value in (
        ("top", CELL_MARGIN_TOP_DXA),
        ("start", CELL_MARGIN_START_DXA),
        ("bottom", CELL_MARGIN_BOTTOM_DXA),
        ("end", CELL_MARGIN_END_DXA),
    ):
        element = OxmlElement(f"w:{tag}")
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")
        margins.append(element)
    _replace_children(table._tbl.tblPr, "w:tblCellMar", margins)


def _column_widths(level_count: int) -> tuple[int, ...]:
    criterion_width = 1800
    weight_width = 900
    available = CONTENT_WIDTH_DXA - criterion_width - weight_width
    base, remainder = divmod(available, level_count)
    level_widths = tuple(
        base + (1 if index < remainder else 0)
        for index in range(level_count)
    )
    widths = (criterion_width, weight_width, *level_widths)
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise AssertionError("Internal table geometry does not sum to the content width.")
    return widths


def _set_table_geometry(table: object, widths: tuple[int, ...]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    table_width = OxmlElement("w:tblW")
    table_width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    table_width.set(qn("w:type"), "dxa")
    _replace_children(tbl_pr, "w:tblW", table_width)

    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    _replace_children(tbl_pr, "w:tblInd", indent)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    _replace_children(tbl_pr, "w:tblLayout", layout)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), TABLE_BORDER)
        borders.append(border)
    _replace_children(tbl_pr, "w:tblBorders", borders)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)

    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths)):
            table.columns[index].width = Twips(width)
            cell.width = Twips(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            cell_width = OxmlElement("w:tcW")
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")
            _replace_children(tc_pr, "w:tcW", cell_width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_table_cell_margins(table)


def _shade_cell(cell: object, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    _replace_children(cell._tc.get_or_add_tcPr(), "w:shd", shading)


def _format_cell(
    cell: object,
    text: str,
    *,
    bold: bool = False,
    centered: bool = False,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    _set_run_font(run, 9.5, bold=bold)


def _mark_repeat_header(row: object) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _canonicalize_docx(raw: bytes) -> bytes:
    """Normalize member order, timestamps, compression, and permissions."""

    source = io.BytesIO(raw)
    target = io.BytesIO()
    with ZipFile(source) as archive:
        members = {
            info.filename: archive.read(info.filename)
            for info in archive.infolist()
            if not info.is_dir()
        }
    with ZipFile(
        target,
        "w",
        compression=ZIP_DEFLATED,
        compresslevel=9,
    ) as archive:
        for name in sorted(members):
            info = ZipInfo(name, date_time=FIXED_PACKAGE_TIMESTAMP)
            info.compress_type = ZIP_DEFLATED
            info.create_system = 3
            info.external_attr = 0o100644 << 16
            archive.writestr(
                info,
                members[name],
                compress_type=ZIP_DEFLATED,
                compresslevel=9,
            )
    return target.getvalue()


def render_docx(spec: TemplateSpec = DEFAULT_SPEC) -> bytes:
    """Render the accepted table-first Word intake shape."""

    _validate_spec(spec)
    document = Document()
    section = document.sections[0]
    section.page_width = Twips(PAGE_WIDTH_DXA)
    section.page_height = Twips(PAGE_HEIGHT_DXA)
    section.top_margin = Twips(MARGIN_DXA)
    section.right_margin = Twips(MARGIN_DXA)
    section.bottom_margin = Twips(MARGIN_DXA)
    section.left_margin = Twips(MARGIN_DXA)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    properties = document.core_properties
    properties.title = "Rubric Weave Intake Template"
    properties.subject = "Synthetic editable rubric intake"
    properties.author = "CourseCraft Workbench"
    properties.last_modified_by = "CourseCraft Workbench"
    properties.category = "Rubric Weave"
    properties.keywords = "rubric, synthetic, template, Brightspace"
    properties.comments = (
        "Generated deterministically over the accepted Workbench producer "
        f"{ACCEPTED_PRODUCER_COMMIT}."
    )
    properties.created = FIXED_DOCUMENT_TIME
    properties.modified = FIXED_DOCUMENT_TIME
    properties.revision = 1

    normal = document.styles["Normal"]
    _set_style_font(normal, 11, "000000")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading = document.styles["Heading 1"]
    _set_style_font(heading, 16, HEADING_BLUE)
    heading.font.bold = True
    heading.paragraph_format.space_before = Pt(DOCX_TITLE_BEFORE_OVERRIDE_PT)
    heading.paragraph_format.space_after = Pt(10)
    heading.paragraph_format.line_spacing = 1.0
    heading.paragraph_format.keep_with_next = True

    title = document.add_paragraph(style="Heading 1")
    title.paragraph_format.keep_with_next = True
    title_run = title.add_run(spec.title)
    _set_run_font(title_run, 16, bold=True)
    title_run.font.color.rgb = RGBColor.from_string(HEADING_BLUE)

    table = document.add_table(
        rows=len(spec.criteria) + 1,
        cols=len(spec.levels) + 2,
    )
    headers = [
        "Criterion",
        "Weight",
        *(
            f"{level.name} ({_format_number(level.score)})"
            for level in spec.levels
        ),
    ]
    for index, (cell, value) in enumerate(zip(table.rows[0].cells, headers)):
        _format_cell(cell, value, bold=True, centered=index != 0)
        _shade_cell(cell, HEADER_FILL)
    _mark_repeat_header(table.rows[0])

    for row, criterion in zip(table.rows[1:], spec.criteria):
        values: Iterable[str] = (
            criterion.name,
            _format_number(criterion.weight),
            *criterion.descriptions,
        )
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            _format_cell(
                cell,
                value,
                bold=index == 0,
                centered=index == 1,
            )

    _set_table_geometry(table, _column_widths(len(spec.levels)))

    for instruction_index, (label, text) in enumerate(INSTRUCTION_PARAGRAPHS):
        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.space_before = (
            Pt(8) if instruction_index == 0 else Pt(0)
        )
        paragraph.paragraph_format.keep_together = True
        label_run = paragraph.add_run(f"{label} ")
        _set_run_font(label_run, 11, bold=True)
        text_run = paragraph.add_run(text)
        _set_run_font(text_run, 11)

    raw = io.BytesIO()
    document.save(raw)
    return _canonicalize_docx(raw.getvalue())


def _sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _manifest(template_payloads: dict[str, bytes]) -> bytes:
    media_types = {
        DOCX_NAME: DOCX_MEDIA_TYPE,
        MARKDOWN_NAME: MARKDOWN_MEDIA_TYPE,
    }
    manifest = {
        "schema": "coursecraft.rubric_weave_template_manifest/1",
        "template_set": TEMPLATE_SET,
        "version": TEMPLATE_VERSION,
        "path_base": "manifest_directory",
        "accepted_producer": {
            "repository": "coursecraft_workbench",
            "commit": ACCEPTED_PRODUCER_COMMIT,
            "authoring_contract": AUTHORING_CONTRACT,
        },
        "templates": [
            {
                "version": TEMPLATE_VERSION,
                "media_type": media_types[path],
                "path": path,
                "bytes": len(template_payloads[path]),
                "sha256": _sha256(template_payloads[path]),
            }
            for path in (DOCX_NAME, MARKDOWN_NAME)
        ],
        "boundaries": {
            "scoring": (
                "Scores and weights are never silently invented; missing or "
                "ambiguous evidence refuses unless the operator explicitly "
                "approves a named fallback."
            ),
            "brightspace_import": (
                "A successful local build or validation is not a Brightspace "
                "import or proof of Brightspace acceptance."
            ),
            "activity_attachment": (
                "The output is rubric-only; attachment to an assignment, "
                "discussion, quiz, or grade item is a separate manual step."
            ),
        },
    }
    return (
        json.dumps(manifest, indent=2, ensure_ascii=False, sort_keys=True) + "\n"
    ).encode("utf-8")


def generated_assets(spec: TemplateSpec = DEFAULT_SPEC) -> dict[str, bytes]:
    """Return every generated canonical payload without touching the filesystem."""

    template_payloads = {
        DOCX_NAME: render_docx(spec),
        MARKDOWN_NAME: render_markdown(spec),
    }
    return {
        **template_payloads,
        MANIFEST_NAME: _manifest(template_payloads),
    }


def write_assets(
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    spec: TemplateSpec = DEFAULT_SPEC,
) -> dict[str, Path]:
    """Write all generated assets and return their exact paths."""

    output_dir.mkdir(parents=True, exist_ok=True)
    result: dict[str, Path] = {}
    for name, data in generated_assets(spec).items():
        path = output_dir / name
        path.write_bytes(data)
        result[name] = path
    return result


def check_assets(output_dir: Path = DEFAULT_OUTPUT_DIR) -> list[str]:
    """Return deterministic-regeneration mismatches without writing files."""

    mismatches: list[str] = []
    for name, expected in generated_assets().items():
        path = output_dir / name
        if not path.is_file():
            mismatches.append(f"missing: {path}")
        elif path.read_bytes() != expected:
            mismatches.append(f"byte mismatch: {path}")
    return mismatches


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate or verify the versioned Rubric Weave intake templates."
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=DEFAULT_OUTPUT_DIR,
        help="Destination directory (defaults to the canonical v1 template lane).",
    )
    parser.add_argument(
        "--check",
        action="store_true",
        help="Refuse if committed assets differ from deterministic regeneration.",
    )
    args = parser.parse_args()

    if args.check:
        mismatches = check_assets(args.output_dir)
        if mismatches:
            for mismatch in mismatches:
                print(mismatch)
            raise SystemExit(1)
        print("VALID")
        return

    for name, path in write_assets(args.output_dir).items():
        data = path.read_bytes()
        print(f"{name}\t{len(data)}\t{_sha256(data)}")


if __name__ == "__main__":
    main()

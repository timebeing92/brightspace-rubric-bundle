#!/usr/bin/env python3
"""Strict Workbench-owned rubric authoring and package orchestration.

This module adapts supported authoring sources into
``coursecraft.rubric_authoring/1``.  It deliberately delegates Brightspace XML
and ZIP construction to :mod:`rubric_package_lib`; there is one package writer.
"""

from __future__ import annotations

import hashlib
import html
import json
import math
import re
import shutil
import subprocess
import tempfile
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from zipfile import BadZipFile, ZipFile

from jsonschema import Draft7Validator

from extract_rubrics_to_workbook import clean_xml_text, rubrics_to_records
from rubric_package_lib import (
    assert_safe_replace_target,
    build_manifest_xml,
    build_orgunit_xml,
    build_rubrics_xml,
    merge_context,
    validate_package_path,
    write_json,
    write_xml,
    zip_package,
)

AUTHORING_SCHEMA = "coursecraft.rubric_authoring/1"
RUN_SCHEMA = "coursecraft.run/1"
SCHEMA_PATH = (
    Path(__file__).resolve().parents[1]
    / "workspace/reference/schemas/rubrics/rubric_authoring_schema.json"
)
RUN_SCHEMA_PATH = (
    Path(__file__).resolve().parents[1]
    / "workspace/reference/schemas/course/run_identity_schema.json"
)
RUBRICS_SCHEMA_PATH = (
    Path(__file__).resolve().parents[1]
    / "workspace/reference/schemas/rubrics/rubrics_schema.json"
)
EXPECTED_PACKAGE_MEMBERS = (
    "imsmanifest.xml",
    "orgunitconfig/orgunitconfig.xml",
    "rubrics_d2l.xml",
)
AUXILIARY_COLUMN_KEYS = {
    "",
    "number",
    "no",
    "row",
    "row_number",
    "id",
    "identifier",
    "index",
    "sort",
    "sort_order",
    "sequence",
    "item",
    "item_number",
    "notes",
    "note",
    "comments",
    "comment",
    "remarks",
    "remark",
    "feedback_notes",
    "evidence",
    "example",
    "examples",
}
MAX_TEXT_SOURCE_BYTES = 10 * 1024 * 1024
MAX_DOCX_SOURCE_BYTES = 20 * 1024 * 1024
MAX_DOCX_MEMBERS = 500
MAX_DOCX_MEMBER_BYTES = 20 * 1024 * 1024
MAX_DOCX_TOTAL_BYTES = 100 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 200


class AuthoringRefusal(ValueError):
    """A policy refusal with structured, machine-readable diagnostics."""

    def __init__(self, diagnostics: list[dict[str, Any]], message: str = "Rubric authoring preflight refused the source."):
        super().__init__(message)
        self.diagnostics = diagnostics


class DuplicateJsonKey(ValueError):
    pass


@dataclass
class DiagnosticBag:
    items: list[dict[str, Any]]

    def add(
        self,
        code: str,
        severity: str,
        message: str,
        location: str,
        remediation: str,
        **extensions: Any,
    ) -> None:
        self.items.append(
            {
                "id": f"diag-{len(self.items) + 1:04d}",
                "code": code,
                "severity": severity,
                "message": message,
                "location": location,
                "remediation": remediation,
                "extensions": extensions,
            }
        )

    @property
    def has_errors(self) -> bool:
        return any(item["severity"] == "error" for item in self.items)


def _preflight_docx_archive(path: Path, bag: DiagnosticBag) -> bool:
    try:
        with ZipFile(path) as archive:
            infos = archive.infolist()
            normalized_names = [info.filename.replace("\\", "/") for info in infos]
            if len(infos) > MAX_DOCX_MEMBERS:
                bag.add(
                    "DOCX_ARCHIVE_UNSAFE",
                    "error",
                    "The DOCX archive contains too many members.",
                    "source",
                    "Provide a smaller, straightforward Word document.",
                    member_count=len(infos),
                    limit=MAX_DOCX_MEMBERS,
                )
            if len(normalized_names) != len(set(normalized_names)):
                bag.add(
                    "DOCX_ARCHIVE_UNSAFE",
                    "error",
                    "The DOCX archive contains duplicate member names.",
                    "source",
                    "Re-save the document in Word before retrying.",
                )
            total = 0
            for info, member in zip(infos, normalized_names):
                parts = Path(member).parts
                mode = (info.external_attr >> 16) & 0o170000
                total += info.file_size
                unsafe_name = (
                    member.startswith("/")
                    or bool(re.match(r"^[A-Za-z]:/", member))
                    or ".." in parts
                )
                if unsafe_name or mode == 0o120000 or info.flag_bits & 0x1:
                    bag.add(
                        "DOCX_ARCHIVE_UNSAFE",
                        "error",
                        "The DOCX archive contains an unsafe member.",
                        "source",
                        "Re-save the document in Word before retrying.",
                    )
                    break
                ratio = info.file_size / max(info.compress_size, 1)
                if (
                    info.file_size > MAX_DOCX_MEMBER_BYTES
                    or (
                        info.file_size > 1024 * 1024
                        and ratio > MAX_DOCX_COMPRESSION_RATIO
                    )
                ):
                    bag.add(
                        "DOCX_ARCHIVE_UNSAFE",
                        "error",
                        "A DOCX archive member exceeds producer safety limits.",
                        "source",
                        "Reduce embedded content and re-save the document.",
                    )
                    break
            if total > MAX_DOCX_TOTAL_BYTES:
                bag.add(
                    "DOCX_ARCHIVE_UNSAFE",
                    "error",
                    "The DOCX archive expands beyond the producer safety limit.",
                    "source",
                    "Reduce embedded content and re-save the document.",
                    uncompressed_bytes=total,
                    limit=MAX_DOCX_TOTAL_BYTES,
                )
            if "word/document.xml" not in normalized_names:
                bag.add(
                    "DOCX_ARCHIVE_UNSAFE",
                    "error",
                    "The DOCX archive is missing word/document.xml.",
                    "source",
                    "Provide a valid Word DOCX file.",
                )
    except (BadZipFile, OSError):
        bag.add(
            "DOCX_ARCHIVE_UNSAFE",
            "error",
            "The DOCX source is not a readable Word archive.",
            "source",
            "Provide a valid Word DOCX file.",
        )
    return not bag.has_errors


def _json_bytes(value: Any) -> bytes:
    return (json.dumps(value, indent=2, ensure_ascii=False, sort_keys=True) + "\n").encode("utf-8")


def _sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def _sha256_file(path: Path) -> str:
    return _sha256_bytes(path.read_bytes())


def _safe_token(value: str, fallback: str) -> str:
    token = re.sub(r"[^A-Za-z0-9_-]+", "_", value.strip())
    token = re.sub(r"_+", "_", token).strip("_-")
    return token or fallback


def _generic_source_label(explicit_label: str | None = None) -> str:
    if explicit_label:
        label = re.sub(r"[/\\\x00-\x1f]+", "_", explicit_label).strip(" .")
        if label:
            return _clean_text(label)
    return "rubric-source"


def _clean_text(value: Any) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value).strip())


def _paragraph_html(value: Any) -> str:
    text = str(value).strip()
    if not text:
        return ""
    if text.startswith("<") and text.endswith(">"):
        return text
    return f"<p>{html.escape(text)}</p>"


def _number(value: Any) -> float | None:
    if isinstance(value, bool) or value is None:
        return None
    if isinstance(value, (int, float)):
        result = float(value)
        return result if math.isfinite(result) else None
    match = re.fullmatch(r"\s*(-?\d+(?:\.\d+)?)\s*%?\s*", str(value))
    if not match:
        return None
    result = float(match.group(1))
    return result if math.isfinite(result) else None


def _header_number(value: str) -> float | None:
    text = _clean_text(value)
    matches = re.findall(r"-?\d+(?:\.\d+)?", text)
    if len(matches) != 1:
        return None
    if len(matches) == 1 and re.fullmatch(r"(?i)level\s+\d+", text):
        return None
    return float(matches[-1])


def _header_number_status(value: str) -> str:
    text = _clean_text(value)
    matches = re.findall(r"-?\d+(?:\.\d+)?", text)
    if not matches or (len(matches) == 1 and re.fullmatch(r"(?i)level\s+\d+", text)):
        return "absent"
    return "valid" if len(matches) == 1 else "ambiguous"


def _clean_level_name(value: str, index: int) -> str:
    text = _clean_text(value)
    text = re.sub(r"\s*\(\s*-?\d+(?:\.\d+)?\s*%?\s*\)\s*$", "", text)
    text = re.sub(r"\s*[-–—]\s*-?\d+(?:\.\d+)?\s*%?\s*$", "", text)
    return text


def _metadata_pairs(value: Any) -> dict[str, float] | None:
    if isinstance(value, dict):
        result: dict[str, float] = {}
        for name, raw in value.items():
            number = _number(raw)
            if number is None:
                return None
            clean_name = _clean_text(name)
            if _identity_key(clean_name) in {_identity_key(item) for item in result}:
                return None
            result[clean_name] = number
        return result
    if isinstance(value, list):
        result = {}
        for item in value:
            if not isinstance(item, dict):
                return None
            name = _clean_text(item.get("name", ""))
            number = _number(
                item.get(
                    "range_start_value",
                    item.get("multiplier", item.get("value", item.get("score"))),
                )
            )
            if (
                not name
                or number is None
                or _identity_key(name) in {_identity_key(existing) for existing in result}
            ):
                return None
            result[name] = number
        return result
    if isinstance(value, str):
        result = {}
        parts = [part.strip() for part in re.split(r"[;,\n]+", value) if part.strip()]
        for part in parts:
            match = re.fullmatch(r"(.+?)\s*[:=]\s*(-?\d+(?:\.\d+)?)\s*%?", part)
            if not match:
                return None
            name = _clean_text(match.group(1))
            if _identity_key(name) in {_identity_key(existing) for existing in result}:
                return None
            result[name] = float(match.group(2))
        return result or None
    return None


def _metadata_pair_names(value: Any) -> list[str]:
    if isinstance(value, dict):
        return [_clean_text(name) for name in value]
    if isinstance(value, list):
        return [
            _clean_text(item.get("name"))
            for item in value
            if isinstance(item, dict)
        ]
    if isinstance(value, str):
        names: list[str] = []
        for part in [item.strip() for item in re.split(r"[;,\n]+", value) if item.strip()]:
            match = re.fullmatch(r"(.+?)\s*[:=].*", part)
            if match:
                names.append(_clean_text(match.group(1)))
        return names
    return []


def _join_metadata_to_levels(
    score_map: dict[str, float],
    level_names: list[str],
) -> dict[str, float] | None:
    """Join named metadata to authored levels by normalized semantic identity."""

    levels_by_identity = {_identity_key(name): name for name in level_names}
    scores_by_identity = {_identity_key(name): value for name, value in score_map.items()}
    if (
        len(levels_by_identity) != len(level_names)
        or len(scores_by_identity) != len(score_map)
        or set(scores_by_identity) != set(levels_by_identity)
    ):
        return None
    return {
        level_name: scores_by_identity[_identity_key(level_name)]
        for level_name in level_names
    }


def _metadata_key(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", value.strip().lower()).strip("_")


def _identity_key(value: str) -> str:
    normalized = unicodedata.normalize("NFKC", _clean_text(value)).casefold()
    return " ".join(normalized.split())


def _strict_json_object(pairs: list[tuple[str, Any]]) -> dict[str, Any]:
    result: dict[str, Any] = {}
    for key, value in pairs:
        if key in result:
            raise DuplicateJsonKey("duplicate JSON object key")
        result[key] = value
    return result


def _reject_json_constant(_value: str) -> None:
    raise ValueError("non-finite JSON number")


def _unknown_fields(
    value: Any,
    allowed: set[str],
) -> set[str]:
    return set(value) - allowed if isinstance(value, dict) else set()


def _split_pipe_row(value: str) -> list[str]:
    text = value.strip()
    if text.startswith("|"):
        text = text[1:]
    if text.endswith("|") and not text.endswith(r"\|"):
        text = text[:-1]
    cells: list[str] = []
    current: list[str] = []
    escaped = False
    for char in text:
        if escaped:
            if char == "|":
                current.append("|")
            else:
                current.extend(["\\", char])
            escaped = False
        elif char == "\\":
            escaped = True
        elif char == "|":
            cells.append("".join(current).strip())
            current = []
        else:
            current.append(char)
    if escaped:
        current.append("\\")
    cells.append("".join(current).strip())
    return cells


def _parse_markdown(path: Path, bag: DiagnosticBag) -> list[dict[str, Any]]:
    lines = path.read_text(encoding="utf-8").splitlines()
    starts = [index for index, line in enumerate(lines) if re.match(r"^##\s+\S", line)]
    if not starts:
        bag.add(
            "SOURCE_SCHEMA_UNSUPPORTED",
            "error",
            "No rubric sections were found.",
            "source",
            "Use a level-two heading for each rubric followed by one pipe table.",
        )
        return []

    rubrics: list[dict[str, Any]] = []
    for section_index, start in enumerate(starts, start=1):
        end = starts[section_index] if section_index < len(starts) else len(lines)
        title = _clean_text(re.sub(r"^##\s+", "", lines[start]))
        block = lines[start + 1 : end]
        metadata: dict[str, str] = {}
        tables: list[list[str]] = []
        current: list[str] = []
        decorative = False
        for line in block:
            stripped = line.strip()
            meta = re.match(r"^-\s*([^:]+):\s*(.+)$", stripped)
            if meta and not current:
                metadata_key = _metadata_key(meta.group(1))
                if metadata_key in metadata:
                    bag.add(
                        "SCORING_METADATA_REQUIRED",
                        "error",
                        "The Markdown rubric repeats a semantic metadata key.",
                        f"rubric[{section_index}].metadata",
                        "Keep exactly one declaration for each metadata key.",
                    )
                else:
                    metadata[metadata_key] = meta.group(2).strip()
            elif stripped.startswith("|"):
                current.append(stripped)
            else:
                if current:
                    tables.append(current)
                    current = []
                if stripped and not stripped.startswith("#"):
                    decorative = True
        if current:
            tables.append(current)
        location = f"rubric[{section_index}]"
        if decorative:
            bag.add(
                "DECORATIVE_CONTENT_IGNORED",
                "info",
                "Narrative material outside the rubric table was ignored.",
                location,
                "Review the normalized contract to confirm only intended rubric content was admitted.",
            )
        if len(tables) != 1 or len(tables[0]) < 3:
            bag.add(
                "SOURCE_SCHEMA_UNSUPPORTED",
                "error",
                "Each Markdown rubric section must contain exactly one usable pipe table.",
                location,
                "Keep one header row, one separator row, and at least one complete criterion row per rubric.",
            )
            continue
        table = tables[0]
        header = _split_pipe_row(table[0])
        separator = _split_pipe_row(table[1])
        if len(separator) != len(header) or any(
            not re.fullmatch(r":?-{3,}:?", cell.replace(" ", ""))
            for cell in separator
        ):
            bag.add(
                "SOURCE_SCHEMA_UNSUPPORTED",
                "error",
                "The Markdown table separator row is missing or malformed.",
                f"{location}.separator",
                "Use one separator cell of at least three hyphens beneath every header.",
            )
            continue
        rows = [_split_pipe_row(row) for row in table[2:]]
        if any(len(row) != len(header) for row in rows):
            bag.add(
                "SOURCE_SCHEMA_UNSUPPORTED",
                "error",
                "The Markdown table contains a non-rectangular row.",
                location,
                "Give every criterion row exactly the same number of cells as the header.",
            )
            continue
        normalized_headers = [_metadata_key(cell) for cell in header]
        auxiliary = [
            header[index]
            for index, value in enumerate(normalized_headers)
            if value in AUXILIARY_COLUMN_KEYS or header[index].strip() == "#"
        ]
        if auxiliary:
            bag.add(
                "SOURCE_SCHEMA_UNSUPPORTED",
                "error",
                "The Markdown rubric contains ambiguous auxiliary columns.",
                f"{location}.header",
                "Remove row-number, ID, notes, comments, and other non-rubric columns.",
                column_count=len(auxiliary),
            )
            continue
        criterion_indexes = [
            index
            for index, value in enumerate(normalized_headers)
            if value in {"criterion", "criteria", "trait", "dimension", "category"}
        ]
        weight_indexes = [
            index
            for index, value in enumerate(normalized_headers)
            if value in {"weight", "points", "point", "pts", "max_points"}
        ]
        if len(criterion_indexes) != 1 or len(weight_indexes) > 1:
            bag.add(
                "DOCX_TABLE_AMBIGUOUS",
                "error",
                "The rubric table does not have one criterion column and at most one weight column.",
                f"{location}.header",
                "Use exactly one Criterion column and zero or one Weight/Points column.",
            )
            continue
        reserved = set(criterion_indexes + weight_indexes)
        level_indexes = [index for index in range(len(header)) if index not in reserved]
        if len(level_indexes) < 2:
            bag.add(
                "SOURCE_SCHEMA_UNSUPPORTED",
                "error",
                "The rubric table has fewer than two performance-level columns.",
                f"{location}.header",
                "Provide at least two ordered performance-level columns.",
            )
            continue
        level_names = [_clean_level_name(header[index], order) for order, index in enumerate(level_indexes, start=1)]
        raw_levels = [
            {
                "name": name,
                "header_number": _header_number(header[column]),
                "header_number_status": _header_number_status(header[column]),
            }
            for name, column in zip(level_names, level_indexes)
        ]
        score_keys = [
            key
            for key in ("level_scores", "level_values", "performance_levels")
            if key in metadata
        ]
        overall_keys = [
            key
            for key in ("overall_thresholds", "overall_levels", "score_bands")
            if key in metadata
        ]
        if len(score_keys) > 1 or len(overall_keys) > 1:
            bag.add(
                "SCORING_METADATA_REQUIRED",
                "error",
                "Competing scoring metadata declarations are ambiguous.",
                f"{location}.metadata",
                "Use exactly one level-score form and one overall-threshold form.",
            )
        level_scores = metadata.get("level_scores") or metadata.get("level_values") or metadata.get("performance_levels")
        score_metadata_status = "absent"
        if level_scores:
            score_metadata_status = "invalid"
            score_map = _metadata_pairs(level_scores)
            if score_map is None:
                bag.add(
                    "SCORING_METADATA_REQUIRED",
                    "error",
                    "Level score metadata could not be parsed.",
                    f"{location}.level_scores",
                    "Use named pairs such as Excellent=100; Capable=75; Beginning=0.",
                )
            else:
                joined_scores = _join_metadata_to_levels(score_map, level_names)
                if joined_scores is None:
                    bag.add(
                        "SCORING_METADATA_REQUIRED",
                        "error",
                        "Level score metadata does not map exactly to the authored levels.",
                        f"{location}.level_scores",
                        "Name every authored level exactly once and include no extra level names.",
                    )
                else:
                    score_metadata_status = "valid"
                    for level in raw_levels:
                        level["score"] = joined_scores[level["name"]]
        overall = (
            metadata.get("overall_thresholds")
            or metadata.get("overall_levels")
            or metadata.get("score_bands")
        )
        criteria = []
        for row_index, row in enumerate(rows, start=1):
            criteria.append(
                {
                    "name": row[criterion_indexes[0]],
                    "weight": row[weight_indexes[0]] if weight_indexes else None,
                    "descriptions": {
                        name: row[column]
                        for name, column in zip(level_names, level_indexes)
                    },
                    "_location": f"{location}.row[{row_index}]",
                }
            )
        rubrics.append(
            {
                "name": title,
                "kind": metadata.get("kind", "other"),
                "source_reference": metadata.get("source_reference"),
                "levels": raw_levels,
                "_level_score_metadata_status": score_metadata_status,
                "overall_levels": overall,
                "criteria": criteria,
                "_location": location,
            }
        )
    return rubrics


def _docx_has_merged_cell(table: Any) -> bool:
    for cell in table._tbl.xpath(".//w:tc"):
        if cell.xpath("./w:tcPr/w:gridSpan"):
            return True
        if cell.xpath("./w:tcPr/w:vMerge"):
            return True
    return False


def _parse_docx(path: Path, bag: DiagnosticBag) -> list[dict[str, Any]]:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = Document(path)
    rubrics: list[dict[str, Any]] = []
    pending: list[str] = []
    table_index = 0
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            text = _clean_text(Paragraph(child, document).text)
            if text:
                pending.append(text)
            continue
        if child.tag != qn("w:tbl"):
            continue
        table_index += 1
        table = Table(child, document)
        location = f"table[{table_index}]"
        if any(cell.tables for row in table.rows for cell in row.cells):
            bag.add(
                "DOCX_NESTED_TABLE_UNSUPPORTED",
                "error",
                "The Word rubric contains a nested table.",
                location,
                "Flatten the rubric to one straightforward table or use Markdown/JSON.",
            )
            pending = []
            continue
        if _docx_has_merged_cell(table):
            bag.add(
                "DOCX_MERGED_CELL_UNSUPPORTED",
                "error",
                "The Word rubric contains a merged cell.",
                location,
                "Unmerge every cell and repeat content explicitly, or use Markdown/JSON.",
            )
            pending = []
            continue
        grid_columns = len(table.columns)
        if grid_columns < 3 or any(len(row.cells) != grid_columns for row in table.rows):
            bag.add(
                "DOCX_TABLE_AMBIGUOUS",
                "error",
                "The Word rubric table is non-rectangular or too narrow.",
                location,
                "Use one rectangular table with Criterion, optional Weight, and at least two levels.",
            )
            pending = []
            continue
        metadata: dict[str, str] = {}
        title_candidates: list[str] = []
        for value in pending:
            match = re.fullmatch(r"([^:]+):\s*(.+)", value)
            if match and _metadata_key(match.group(1)) in {
                "level_scores",
                "level_values",
                "performance_levels",
                "overall_thresholds",
                "overall_levels",
                "score_bands",
                "kind",
                "source_reference",
            }:
                metadata_key = _metadata_key(match.group(1))
                if metadata_key in metadata:
                    bag.add(
                        "DOCX_TABLE_AMBIGUOUS",
                        "error",
                        "The Word table has repeated semantic metadata.",
                        f"{location}.metadata",
                        "Keep exactly one declaration for each metadata key.",
                    )
                else:
                    metadata[metadata_key] = match.group(2).strip()
            else:
                title_candidates.append(value)
        pending = []
        if len(title_candidates) != 1:
            bag.add(
                "DOCX_TABLE_AMBIGUOUS",
                "error",
                "The Word table does not have exactly one unambiguous preceding rubric title.",
                location,
                "Place one plain or heading paragraph immediately before the table; label metadata as Key: Value.",
                title_candidate_count=len(title_candidates),
            )
            continue
        values = [[_clean_text(cell.text) for cell in row.cells] for row in table.rows]
        if len(values) < 2:
            bag.add(
                "DOCX_TABLE_AMBIGUOUS",
                "error",
                "The Word rubric table has no criterion rows.",
                location,
                "Add at least one complete criterion row.",
            )
            continue
        # Reuse the strict Markdown table parser by materializing an in-memory
        # section-shaped representation through the same column rules.
        header = values[0]
        normalized_headers = [_metadata_key(value) for value in header]
        auxiliary = [
            header[index]
            for index, value in enumerate(normalized_headers)
            if value in AUXILIARY_COLUMN_KEYS or header[index].strip() == "#"
        ]
        if auxiliary:
            bag.add(
                "DOCX_TABLE_AMBIGUOUS",
                "error",
                "The Word rubric contains ambiguous auxiliary columns.",
                f"{location}.header",
                "Remove row-number, ID, notes, comments, and other non-rubric columns.",
                column_count=len(auxiliary),
            )
            continue
        criterion_indexes = [
            index for index, value in enumerate(normalized_headers)
            if value in {"criterion", "criteria", "trait", "dimension", "category"}
        ]
        weight_indexes = [
            index for index, value in enumerate(normalized_headers)
            if value in {"weight", "points", "point", "pts", "max_points"}
        ]
        if len(criterion_indexes) != 1 or len(weight_indexes) > 1:
            bag.add(
                "DOCX_TABLE_AMBIGUOUS",
                "error",
                "The Word table does not have one criterion column and at most one weight column.",
                f"{location}.header",
                "Use unique Criterion and Weight/Points headers.",
            )
            continue
        reserved = set(criterion_indexes + weight_indexes)
        level_indexes = [index for index in range(len(header)) if index not in reserved]
        if len(level_indexes) < 2:
            bag.add(
                "DOCX_TABLE_AMBIGUOUS",
                "error",
                "The Word table has fewer than two performance-level columns.",
                f"{location}.header",
                "Provide at least two uniquely named level columns.",
            )
            continue
        level_names = [_clean_level_name(header[index], order) for order, index in enumerate(level_indexes, start=1)]
        if len({_metadata_key(name) for name in level_names}) != len(level_names):
            bag.add(
                "DOCX_TABLE_AMBIGUOUS",
                "error",
                "The Word table has duplicate semantic performance-level headers.",
                f"{location}.header",
                "Give every level a unique name.",
            )
            continue
        levels = [
            {
                "name": name,
                "header_number": _header_number(header[column]),
                "header_number_status": _header_number_status(header[column]),
            }
            for name, column in zip(level_names, level_indexes)
        ]
        score_keys = [
            key
            for key in ("level_scores", "level_values", "performance_levels")
            if key in metadata
        ]
        overall_keys = [
            key
            for key in ("overall_thresholds", "overall_levels", "score_bands")
            if key in metadata
        ]
        if len(score_keys) > 1 or len(overall_keys) > 1:
            bag.add(
                "DOCX_TABLE_AMBIGUOUS",
                "error",
                "Competing Word scoring metadata declarations are ambiguous.",
                f"{location}.metadata",
                "Use exactly one level-score form and one overall-threshold form.",
            )
        raw_score_metadata = (
            metadata.get("level_scores")
            or metadata.get("level_values")
            or metadata.get("performance_levels")
        )
        score_metadata_status = "absent"
        if raw_score_metadata:
            score_metadata_status = "invalid"
            score_map = _metadata_pairs(raw_score_metadata)
            if score_map is None:
                bag.add(
                    "SCORING_METADATA_REQUIRED",
                    "error",
                    "Word level-score metadata is invalid or duplicated.",
                    f"{location}.level_scores",
                    "Use exactly one named numeric score for every level.",
                )
            else:
                joined_scores = _join_metadata_to_levels(score_map, level_names)
                if joined_scores is None:
                    bag.add(
                        "SCORING_METADATA_REQUIRED",
                        "error",
                        "Word level-score metadata does not map exactly to the authored levels.",
                        f"{location}.level_scores",
                        "Name every authored level exactly once and include no extra level names.",
                    )
                else:
                    score_metadata_status = "valid"
                    for level in levels:
                        level["score"] = joined_scores[level["name"]]
        criteria = []
        for row_index, row in enumerate(values[1:], start=1):
            criteria.append(
                {
                    "name": row[criterion_indexes[0]],
                    "weight": row[weight_indexes[0]] if weight_indexes else None,
                    "descriptions": {
                        name: row[column]
                        for name, column in zip(level_names, level_indexes)
                    },
                    "_location": f"{location}.row[{row_index}]",
                }
            )
        rubrics.append(
            {
                "name": title_candidates[0],
                "kind": metadata.get("kind", "other"),
                "source_reference": metadata.get("source_reference"),
                "levels": levels,
                "_level_score_metadata_status": score_metadata_status,
                "overall_levels": (
                    metadata.get("overall_thresholds")
                    or metadata.get("overall_levels")
                    or metadata.get("score_bands")
                ),
                "criteria": criteria,
                "_location": location,
            }
        )
    if not document.tables:
        bag.add(
            "SOURCE_SCHEMA_UNSUPPORTED",
            "error",
            "The Word document contains no rubric tables.",
            "source",
            "Provide at least one straightforward Word table or use Markdown/JSON.",
        )
    return rubrics


def _expected_extracted_attributes(source: dict[str, Any]) -> dict[str, str]:
    """Exact analytic rubric attributes emitted by the shared Workbench writer."""

    return {
        "id": str(source.get("id", "")),
        "resource_code": str(source.get("resource_code", "")),
        "name": str(source.get("name", "")),
        "type": "1",
        "scoring_method": "3",
        "display_levels_in_des_order": "True",
        "state": "0",
        "visibility": "0",
        "uses_overall_score": "True",
        "has_manual_alignment": "False",
        "score_visible_to_assessed_users": "True",
        "enabled_feedback_copy": "False",
        "usage_restrictions": "Competency,ePortfolio",
    }


def _score_band_matches(name: str, score_band: str, threshold: float) -> bool:
    match = re.fullmatch(r"(-?\d+(?:\.\d+)?)%(\+?)", score_band)
    if not match:
        return False
    expected_plus = "" if name == "Not Demonstrated" else "+"
    parsed = _number(match.group(1))
    return (
        parsed is not None
        and match.group(2) == expected_plus
        and abs(parsed - threshold) <= 1e-8
    )


def _adapt_extracted(data: dict[str, Any], bag: DiagnosticBag) -> list[dict[str, Any]]:
    def reject(message: str, location: str, remediation: str) -> None:
        bag.add(
            "RUBRICS_ADAPTER_UNTRANSLATABLE",
            "error",
            message,
            location,
            remediation,
        )

    if data.get("diagnostics"):
        reject(
            "The extraction contract contains unresolved diagnostics.",
            "source.diagnostics",
            "Resolve the extraction diagnostics before authoring adaptation.",
        )
        return []
    source_rubrics = data.get("rubrics") or []
    if not source_rubrics:
        reject(
            "The extraction contract contains no rubric records.",
            "source.rubrics",
            "Provide a canonical extraction containing at least one rubric.",
        )
        return []

    ids = [str(item.get("id", "")) for item in source_rubrics]
    resource_codes = [str(item.get("resource_code", "")) for item in source_rubrics]
    rubric_names = [str(item.get("name", "")) for item in source_rubrics]
    if (
        any(not value or value != value.strip() for value in ids + resource_codes)
        or any(not value or value != _clean_text(value) for value in rubric_names)
        or len(set(ids)) != len(ids)
        or len({_identity_key(value) for value in resource_codes}) != len(resource_codes)
        or len({_identity_key(value) for value in rubric_names}) != len(rubric_names)
    ):
        reject(
            "Extracted rubric identities are incomplete, duplicated, or unstable.",
            "source.rubrics",
            "Use unique non-empty rubric IDs, resource codes, and names.",
        )
        return []

    result: list[dict[str, Any]] = []
    for rubric_index, source in enumerate(source_rubrics, start=1):
        location = f"rubric[{rubric_index}]"
        rubric_valid = True

        if (
            source.get("scoring_method") != "3"
            or source.get("attributes") != _expected_extracted_attributes(source)
        ):
            reject(
                "The extracted rubric uses unsupported analytic attributes or state.",
                f"{location}.attributes",
                "Adapt only exact analytic rubric state reproducible by the Workbench writer.",
            )
            rubric_valid = False
        if _clean_text(source.get("description", "")):
            reject(
                "The extracted rubric contains description prose the authoring contract cannot preserve.",
                f"{location}.description",
                "Remove the prose in Brightspace or author a new rubric contract.",
            )
            rubric_valid = False

        source_levels = source.get("levels") or []
        names = [str(level.get("name", "")) for level in source_levels]
        level_ids = [str(level.get("level_id", "")) for level in source_levels]
        level_orders = [level.get("sort_order") for level in source_levels]
        if (
            len(names) < 2
            or any(not value or value != _clean_text(value) for value in names)
            or any(not value or value != value.strip() for value in level_ids)
            or len({_identity_key(value) for value in names}) != len(names)
            or len(set(level_ids)) != len(level_ids)
            or any(
                isinstance(value, bool)
                or not isinstance(value, int)
                or value < 1
                for value in level_orders
            )
            or len(set(level_orders)) != len(level_orders)
            or level_orders != sorted(level_orders)
        ):
            reject(
                "Extracted rubric levels are incomplete, duplicated, or unordered.",
                f"{location}.levels",
                "Use non-empty unique level IDs, names, and ordered sort positions.",
            )
            rubric_valid = False

        raw_criteria = source.get("criteria") or []
        criterion_names = [str(item.get("name", "")) for item in raw_criteria]
        criterion_orders = [item.get("sort_order") for item in raw_criteria]
        if (
            not raw_criteria
            or any(
                not value or value != _clean_text(value)
                for value in criterion_names
            )
            or len({_identity_key(value) for value in criterion_names})
            != len(criterion_names)
            or any(
                isinstance(value, bool)
                or not isinstance(value, int)
                or value < 1
                for value in criterion_orders
            )
            or len(set(criterion_orders)) != len(criterion_orders)
            or criterion_orders != sorted(criterion_orders)
        ):
            reject(
                "Extracted criteria are incomplete, duplicated, or unordered.",
                f"{location}.criteria",
                "Use non-empty unique criterion names and ordered sort positions.",
            )
            rubric_valid = False

        multiplier_sequences: list[list[float]] = []
        criteria: list[dict[str, Any]] = []
        maxima: list[float] = []
        if rubric_valid:
            for criterion_index, criterion in enumerate(raw_criteria, start=1):
                criterion_location = f"{location}.criterion[{criterion_index}]"
                cells = criterion.get("cells") or []
                by_id = {str(cell.get("level_id", "")): cell for cell in cells}
                if (
                    len(cells) != len(level_ids)
                    or len(by_id) != len(cells)
                    or set(by_id) != set(level_ids)
                    or any(
                        str(by_id[level_id].get("level_name", "")) != expected_name
                        for level_id, expected_name in zip(level_ids, names)
                    )
                ):
                    reject(
                        "A criterion does not join exactly once to every level ID and name.",
                        criterion_location,
                        "Use a complete bijective extracted criterion grid.",
                    )
                    rubric_valid = False
                    continue
                numeric_points: list[float] = []
                scores_valid = True
                for level_id in level_ids:
                    cell = by_id[level_id]
                    points = _number(cell.get("points"))
                    points_raw = _number(cell.get("points_raw"))
                    if (
                        points is None
                        or points_raw is None
                        or abs(points - points_raw) > 1e-8
                        or not _clean_text(cell.get("description", ""))
                    ):
                        scores_valid = False
                        break
                    numeric_points.append(float(points))
                if not scores_valid:
                    reject(
                        "A criterion contains unsupported scores, raw-score disagreement, or empty descriptions.",
                        criterion_location,
                        "Use finite matching points/points_raw values and complete descriptions.",
                    )
                    rubric_valid = False
                    continue
                maximum = max(numeric_points)
                if maximum <= 0:
                    reject(
                        "A criterion has no positive maximum score.",
                        criterion_location,
                        "Use a positive criterion maximum.",
                    )
                    rubric_valid = False
                    continue
                sequence = [value / maximum for value in numeric_points]
                if (
                    any(value < 0 or value > 1 for value in sequence)
                    or len({round(value, 12) for value in sequence}) != len(sequence)
                    or any(
                        later >= earlier
                        for earlier, later in zip(sequence, sequence[1:])
                    )
                    or abs(sequence[0] - 1.0) > 1e-8
                    or abs(sequence[-1]) > 1e-8
                ):
                    reject(
                        "A criterion yields an unsupported multiplier sequence.",
                        criterion_location,
                        "Use one strictly descending unique multiplier sequence from 1 to 0.",
                    )
                    rubric_valid = False
                    continue
                maxima.append(maximum)
                multiplier_sequences.append(sequence)
                criteria.append(
                    {
                        "name": criterion["name"],
                        "weight": maximum,
                        "weight_source": "extracted_cell_points",
                        "descriptions": {
                            name: by_id[level_id]["description"]
                            for name, level_id in zip(names, level_ids)
                        },
                        "_location": criterion_location,
                    }
                )

        if rubric_valid and abs(sum(maxima) - 100.0) > 1e-8:
            reject(
                "Extracted criterion maxima do not total the 100-point weight basis.",
                f"{location}.criteria",
                "Use analytic criterion maxima totaling exactly 100.",
            )
            rubric_valid = False

        reference = multiplier_sequences[0] if multiplier_sequences else []
        if rubric_valid and any(
            len(sequence) != len(reference)
            or any(abs(left - right) > 1e-8 for left, right in zip(sequence, reference))
            for sequence in multiplier_sequences[1:]
        ):
            reject(
                "Criterion scores do not yield one consistent multiplier sequence.",
                f"{location}.criteria",
                "Author explicit scoring metadata instead of adapting this extraction record.",
            )
            rubric_valid = False

        overall = source.get("overall_levels") or []
        overall_names = [str(item.get("name", "")) for item in overall]
        overall_orders = [item.get("sort_order") for item in overall]
        thresholds = [_number(item.get("range_start_value")) for item in overall]
        if (
            len(overall) != len(source_levels)
            or overall_names != names
            or overall_orders != level_orders
            or len({_identity_key(value) for value in overall_names})
            != len(overall_names)
            or len(set(overall_orders)) != len(overall_orders)
            or any(value is None for value in thresholds)
            or any(
                _clean_text(item.get("description", ""))
                or _clean_text(item.get("feedback", ""))
                for item in overall
            )
        ):
            reject(
                "Overall levels do not exactly match level identity, order, or supported empty prose.",
                f"{location}.overall_levels",
                "Use one empty-prose overall record for every level in the same order.",
            )
            rubric_valid = False

        numeric_thresholds = [
            float(value) for value in thresholds if value is not None
        ]
        if rubric_valid and (
            len({round(value, 9) for value in numeric_thresholds})
            != len(numeric_thresholds)
            or any(value < 0 or value > 100 for value in numeric_thresholds)
            or any(
                abs((threshold / 100.0) - multiplier) > 1e-8
                for threshold, multiplier in zip(numeric_thresholds, reference)
            )
            or any(
                not _score_band_matches(name, level.get("score_band", ""), threshold)
                for name, level, threshold in zip(
                    names, source_levels, numeric_thresholds
                )
            )
        ):
            reject(
                "Overall thresholds or score bands disagree with the cell-derived scale.",
                f"{location}.overall_levels",
                "Use unique finite thresholds and canonical score-band labels matching every level.",
            )
            rubric_valid = False

        if not rubric_valid:
            continue
        overall_map = dict(zip(names, numeric_thresholds))
        result.append(
            {
                "name": source["name"],
                "kind": "other",
                "source_reference": source["resource_code"],
                "source_id": source["id"],
                "resource_code": source["resource_code"],
                "levels": [
                    {
                        "name": name,
                        "multiplier": multiplier,
                        "score_source": "extracted_cell_points",
                    }
                    for name, multiplier in zip(names, reference)
                ],
                "overall_levels": overall_map,
                "overall_source": "extracted_overall_threshold",
                "criteria": criteria,
                "_location": location,
            }
        )
    return result


def _parse_json(path: Path, bag: DiagnosticBag) -> tuple[list[dict[str, Any]], str, dict[str, Any]]:
    try:
        data = json.loads(
            path.read_text(encoding="utf-8"),
            object_pairs_hook=_strict_json_object,
            parse_constant=_reject_json_constant,
        )
    except (OSError, json.JSONDecodeError, DuplicateJsonKey, ValueError) as exc:
        bag.add(
            "SOURCE_SCHEMA_UNSUPPORTED",
            "error",
            "The JSON source is not valid UTF-8 JSON.",
            "source",
            "Provide a valid UTF-8 JSON object.",
            error_type=type(exc).__name__,
        )
        return [], "legacy_builder_json", {}
    if not isinstance(data, dict):
        bag.add(
            "SOURCE_SCHEMA_UNSUPPORTED",
            "error",
            "Top-level JSON is not an object.",
            "source",
            "Use coursecraft.rubric_authoring/1, coursecraft.rubrics/1, or the builder JSON shape.",
        )
        return [], "legacy_builder_json", {}
    if data.get("schema") == "coursecraft.rubrics/1":
        schema = json.loads(RUBRICS_SCHEMA_PATH.read_text(encoding="utf-8"))
        schema_errors = sorted(
            Draft7Validator(schema).iter_errors(data),
            key=lambda item: list(item.path),
        )
        for error in schema_errors:
            bag.add(
                "RUBRICS_ADAPTER_UNTRANSLATABLE",
                "error",
                "The extraction JSON does not satisfy its declared schema.",
                ".".join(str(item) for item in error.path) or "source",
                "Regenerate the extraction JSON with the canonical extractor.",
                validator=error.validator,
            )
        if schema_errors:
            return [], "coursecraft.rubrics/1", data
        return _adapt_extracted(data, bag), "coursecraft.rubrics/1", data
    if data.get("schema") == AUTHORING_SCHEMA:
        schema = json.loads(SCHEMA_PATH.read_text(encoding="utf-8"))
        schema_errors = sorted(
            Draft7Validator(schema).iter_errors(data),
            key=lambda item: list(item.path),
        )
        for error in schema_errors:
            bag.add(
                "SOURCE_SCHEMA_UNSUPPORTED",
                "error",
                "The authoring JSON does not satisfy its declared schema.",
                ".".join(str(item) for item in error.path) or "source",
                "Correct the source contract against coursecraft.rubric_authoring/1.",
                validator=error.validator,
            )
        if schema_errors:
            return [], AUTHORING_SCHEMA, data
        if any(
            isinstance(item, dict) and item.get("severity") == "error"
            for item in data.get("diagnostics", [])
        ):
            bag.add(
                "SOURCE_SCHEMA_UNSUPPORTED",
                "error",
                "The input authoring contract carries unresolved error diagnostics.",
                "source.diagnostics",
                "Resolve the recorded errors before re-ingesting the contract.",
            )
            return [], AUTHORING_SCHEMA, data
        rubrics = []
        for index, rubric in enumerate(data.get("rubrics") or [], start=1):
            rubrics.append(
                {
                    "name": rubric.get("name"),
                    "kind": rubric.get("kind", "other"),
                    "source_reference": rubric.get("source_reference"),
                    "levels": [dict(level) for level in rubric.get("levels") or []],
                    "overall_levels": [dict(value) for value in rubric.get("overall_levels") or []],
                    "overall_source": next(
                        iter(
                            {
                                value.get("source")
                                for value in rubric.get("overall_levels") or []
                                if value.get("source")
                            }
                        ),
                        None,
                    ),
                    "_overall_sources": [
                        value.get("source")
                        for value in rubric.get("overall_levels") or []
                    ],
                    "criteria": [
                        {
                            "name": criterion.get("name"),
                            "weight": criterion.get("weight"),
                            "weight_source": criterion.get("weight_source", "explicit_weight"),
                            "descriptions": dict(criterion.get("descriptions") or {}),
                            "_location": f"rubric[{index}].criterion[{criterion_index}]",
                        }
                        for criterion_index, criterion in enumerate(rubric.get("criteria") or [], start=1)
                    ],
                    "_location": f"rubric[{index}]",
                }
            )
        return rubrics, AUTHORING_SCHEMA, data
    if data.get("schema"):
        bag.add(
            "SOURCE_SCHEMA_UNSUPPORTED",
            "error",
            "The JSON source declares an unsupported schema.",
            "source.schema",
            "Use coursecraft.rubric_authoring/1, coursecraft.rubrics/1, or omit schema for builder JSON.",
        )
        return [], "legacy_builder_json", data
    if not isinstance(data.get("rubrics"), list) or not data["rubrics"]:
        bag.add(
            "SOURCE_SCHEMA_UNSUPPORTED",
            "error",
            "Builder JSON must contain a non-empty rubrics array.",
            "source.rubrics",
            "Add at least one rubric object.",
        )
        return [], "legacy_builder_json", data
    allowed_top = {"package", "rubrics"}
    unknown_top = _unknown_fields(data, allowed_top)
    if unknown_top:
        bag.add(
            "SOURCE_SCHEMA_UNSUPPORTED",
            "error",
            "Legacy builder JSON contains unsupported top-level fields.",
            "source",
            "Remove activity, attachment, association, and unknown fields.",
            field_count=len(unknown_top),
        )
        return [], "legacy_builder_json", data
    allowed_package = {
        "manifest_identifier",
        "resource_prefix",
        "rubric_id_start",
        "resource_code_start",
        "level_id_base_start",
        "orgunit",
    }
    if not isinstance(data.get("package", {}), dict) or _unknown_fields(
        data.get("package", {}), allowed_package
    ):
        bag.add(
            "SOURCE_SCHEMA_UNSUPPORTED",
            "error",
            "Legacy builder package metadata has an invalid shape or unsupported fields.",
            "source.package",
            "Use only documented deterministic package and org-unit fields.",
        )
        return [], "legacy_builder_json", data
    rubrics = []
    for rubric_index, rubric in enumerate(data["rubrics"], start=1):
        allowed_rubric = {
            "name",
            "source_id",
            "source_reference",
            "resource_code",
            "levels",
            "overall_levels",
            "overall_thresholds",
            "criteria",
        }
        if not isinstance(rubric, dict) or _unknown_fields(rubric, allowed_rubric):
            bag.add(
                "SOURCE_SCHEMA_UNSUPPORTED",
                "error",
                "A legacy rubric object has an invalid shape or unsupported fields.",
                f"rubric[{rubric_index}]",
                "Remove activity, attachment, association, and unknown fields.",
            )
            continue
        raw_levels = rubric.get("levels")
        if not raw_levels and rubric.get("criteria"):
            first = rubric["criteria"][0]
            if isinstance(first.get("levels"), dict):
                raw_levels = [{"name": name} for name in first["levels"]]
        levels = []
        for level_index, level in enumerate(raw_levels or [], start=1):
            if isinstance(level, str):
                levels.append({"name": level})
            elif isinstance(level, dict):
                allowed_level = {
                    "name",
                    "multiplier",
                    "value",
                    "score",
                }
                if _unknown_fields(level, allowed_level) or not _clean_text(level.get("name")):
                    bag.add(
                        "SOURCE_SCHEMA_UNSUPPORTED",
                        "error",
                        "A legacy level has an invalid shape or unsupported fields.",
                        f"rubric[{rubric_index}].level[{level_index}]",
                        "Provide a string label or an object with a non-empty authored name.",
                    )
                    continue
                levels.append(dict(level))
            else:
                bag.add(
                    "SOURCE_SCHEMA_UNSUPPORTED",
                    "error",
                    "A legacy level is not a string or object.",
                    f"rubric[{rubric_index}].level[{level_index}]",
                    "Provide a string label or an object with a non-empty authored name.",
                )
        raw_criteria = rubric.get("criteria") or []
        if not isinstance(raw_criteria, list):
            bag.add(
                "SOURCE_SCHEMA_UNSUPPORTED",
                "error",
                "Legacy rubric criteria must be an array.",
                f"rubric[{rubric_index}].criteria",
                "Provide an array of criterion objects.",
            )
            raw_criteria = []
        parsed_criteria: list[dict[str, Any]] = []
        for criterion_index, criterion in enumerate(raw_criteria, start=1):
            allowed_criterion = {
                "name",
                "weight",
                "levels",
                "descriptions",
            }
            if not isinstance(criterion, dict) or _unknown_fields(
                criterion, allowed_criterion
            ):
                bag.add(
                    "SOURCE_SCHEMA_UNSUPPORTED",
                    "error",
                    "A legacy criterion has an invalid shape or unsupported fields.",
                    f"rubric[{rubric_index}].criterion[{criterion_index}]",
                    "Use only name, weight provenance, and level descriptions.",
                )
                continue
            descriptions = criterion.get("descriptions", criterion.get("levels", {}))
            if not isinstance(descriptions, dict):
                bag.add(
                    "SOURCE_SCHEMA_UNSUPPORTED",
                    "error",
                    "Legacy criterion descriptions must be an object.",
                    f"rubric[{rubric_index}].criterion[{criterion_index}]",
                    "Map each authored level name to one description.",
                )
                continue
            parsed_criteria.append(
                {
                    "name": criterion.get("name"),
                    "weight": criterion.get("weight"),
                    "descriptions": dict(descriptions),
                    "_location": f"rubric[{rubric_index}].criterion[{criterion_index}]",
                }
            )
        rubrics.append(
            {
                "name": rubric.get("name"),
                "kind": "other",
                "source_reference": rubric.get("source_reference") or rubric.get("source_id"),
                "source_id": rubric.get("source_id"),
                "resource_code": rubric.get("resource_code"),
                "levels": levels,
                "overall_levels": rubric.get("overall_levels", rubric.get("overall_thresholds")),
                "criteria": parsed_criteria,
                "_location": f"rubric[{rubric_index}]",
            }
        )
    return rubrics, "legacy_builder_json", data


def _normalize_rubric(
    raw: dict[str, Any],
    bag: DiagnosticBag,
    allow_even_spacing: bool,
    allow_equal_weights: bool,
    adapter: str,
) -> dict[str, Any] | None:
    location = raw.get("_location", "rubric")
    name = _clean_text(raw.get("name", ""))
    if not name:
        bag.add(
            "SOURCE_SCHEMA_UNSUPPORTED",
            "error",
            "A rubric has no name.",
            f"{location}.name",
            "Provide a non-empty rubric name.",
        )
        return None
    raw_levels = raw.get("levels") or []
    if len(raw_levels) < 2:
        bag.add(
            "SCORING_METADATA_REQUIRED",
            "error",
            "The rubric does not expose at least two performance levels.",
            f"{location}.levels",
            "Provide two or more ordered level definitions.",
        )
        return None
    missing_level_positions = [
        str(index)
        for index, level in enumerate(raw_levels, start=1)
        if not _clean_text(level if isinstance(level, str) else level.get("name"))
    ]
    if missing_level_positions:
        bag.add(
            "SOURCE_SCHEMA_UNSUPPORTED",
            "error",
            "One or more performance levels have no authored name.",
            f"{location}.levels",
            "Provide a non-empty name for every level.",
            positions=missing_level_positions,
        )
        return None
    level_names = [
        _clean_level_name(
            level if isinstance(level, str) else level.get("name", ""),
            index,
        )
        for index, level in enumerate(raw_levels, start=1)
    ]
    if len({_identity_key(value) for value in level_names}) != len(level_names):
        bag.add(
            "SOURCE_SCHEMA_UNSUPPORTED",
            "error",
            "Performance-level names are duplicated.",
            f"{location}.levels",
            "Give every performance level a unique name.",
        )
        return None
    if raw.get("_level_score_metadata_status") == "invalid":
        return None
    raw_overall = raw.get("overall_levels")
    overall_present = raw_overall not in (None, "", [], {})
    overall_names = _metadata_pair_names(raw_overall)
    parsed_overall = _metadata_pairs(raw_overall)
    overall_invalid = overall_present and (
        parsed_overall is None
        or set(parsed_overall) != set(level_names)
        or len({_identity_key(name) for name in overall_names}) != len(overall_names)
    )
    if overall_invalid:
        bag.add(
            "SCORING_METADATA_REQUIRED",
            "error",
            "Explicit overall-threshold metadata is invalid, duplicated, or incomplete.",
            f"{location}.overall_levels",
            "Provide exactly one numeric threshold for every named level.",
        )
        return None

    explicit: list[float | None] = []
    explicit_declared: list[bool] = []
    declared_level_sources = [
        str(level.get("score_source", ""))
        for level in raw_levels
        if isinstance(level, dict) and level.get("score_source")
    ]
    declared_level_source: str | None = None
    if adapter == AUTHORING_SCHEMA:
        unique_sources = set(declared_level_sources)
        overall_sources = {
            str(value)
            for value in raw.get("_overall_sources", [])
            if value
        }
        if len(declared_level_sources) != len(raw_levels) or len(unique_sources) != 1:
            bag.add(
                "SCORING_METADATA_REQUIRED",
                "error",
                "Declared authoring level provenance is incomplete or mixed.",
                f"{location}.levels",
                "Use one uniform score source for every level.",
            )
            return None
        declared_level_source = next(iter(unique_sources))
        if declared_level_source == "extracted_cell_points":
            bag.add(
                "RUBRICS_ADAPTER_UNTRANSLATABLE",
                "error",
                "Declared authoring JSON cannot self-assert extracted scoring provenance.",
                f"{location}.levels",
                "Re-adapt the original coursecraft.rubrics/1 extraction contract.",
            )
            return None
        if len(overall_sources) != 1:
            bag.add(
                "SCORING_METADATA_REQUIRED",
                "error",
                "Declared authoring overall-threshold provenance is incomplete or mixed.",
                f"{location}.overall_levels",
                "Use one uniform threshold source for every overall level.",
            )
            return None
        if overall_sources == {"extracted_overall_threshold"}:
            bag.add(
                "RUBRICS_ADAPTER_UNTRANSLATABLE",
                "error",
                "Declared authoring JSON cannot self-assert extracted threshold provenance.",
                f"{location}.overall_levels",
                "Re-adapt the original coursecraft.rubrics/1 extraction contract.",
            )
            return None
    elif adapter == "coursecraft.rubrics/1":
        if (
            len(declared_level_sources) != len(raw_levels)
            or set(declared_level_sources) != {"extracted_cell_points"}
        ):
            bag.add(
                "RUBRICS_ADAPTER_UNTRANSLATABLE",
                "error",
                "Extraction scoring provenance is incomplete.",
                f"{location}.levels",
                "Regenerate and re-adapt the extraction contract.",
            )
            return None
        declared_level_source = "extracted_cell_points"
    elif declared_level_sources:
        bag.add(
            "SOURCE_SCHEMA_UNSUPPORTED",
            "error",
            "This source adapter cannot declare scoring provenance.",
            f"{location}.levels",
            "Remove provenance fields and provide only explicit numeric scoring values.",
        )
        return None
    for level in raw_levels:
        if not isinstance(level, dict):
            explicit.append(None)
            explicit_declared.append(False)
            continue
        explicit_declared.append(
            any(key in level for key in ("multiplier", "value", "score"))
        )
        value = level.get("multiplier")
        if value is None:
            value = level.get("value", level.get("score"))
            number = _number(value)
            if number is not None and number > 1:
                number /= 100.0
        else:
            number = _number(value)
        explicit.append(number)
    score_source: str | None = None
    multipliers: list[float] = []
    if any(explicit_declared) and (
        not all(explicit_declared)
        or not all(value is not None for value in explicit)
    ):
        bag.add(
            "SCORING_METADATA_REQUIRED",
            "error",
            "Explicit level scoring metadata is invalid or incomplete.",
            f"{location}.levels",
            "Provide one parseable multiplier, value, or score for every level.",
        )
    elif explicit and all(value is not None for value in explicit):
        multipliers = [float(value) for value in explicit if value is not None]
        score_source = declared_level_source or "explicit_level_metadata"
    else:
        if parsed_overall is not None and set(parsed_overall) == set(level_names):
            if max(parsed_overall.values()) > 0:
                multipliers = [parsed_overall[level_name] / 100.0 for level_name in level_names]
                score_source = "overall_threshold_metadata"
        if score_source is None:
            header_statuses = [
                str(level.get("header_number_status", "absent"))
                if isinstance(level, dict)
                else "absent"
                for level in raw_levels
            ]
            if "ambiguous" in header_statuses or (
                "valid" in header_statuses
                and any(status != "valid" for status in header_statuses)
            ):
                bag.add(
                    "SCORING_METADATA_REQUIRED",
                    "error",
                    "Numeric level-header metadata is ambiguous or incomplete.",
                    f"{location}.levels",
                    "Use exactly one numeric token in every level header or explicit metadata.",
                )
                return None
            header_numbers = [
                _number(level.get("header_number")) if isinstance(level, dict) else None
                for level in raw_levels
            ]
            if header_numbers and all(value is not None for value in header_numbers):
                maximum = max(float(value) for value in header_numbers if value is not None)
                if maximum > 0:
                    multipliers = [
                        float(value) / maximum for value in header_numbers if value is not None
                    ]
                    score_source = "numeric_level_header"
        if score_source is None and allow_even_spacing:
            denominator = len(level_names) - 1
            multipliers = [1.0 - index / denominator for index in range(len(level_names))]
            score_source = "approved_even_spacing"
            bag.add(
                "EVEN_SPACING_APPROVED",
                "warning",
                "Even level spacing was applied under explicit operator approval.",
                f"{location}.levels",
                "Review the generated multipliers and mapping before import.",
            )
    if score_source is None:
        bag.add(
            "SCORING_METADATA_REQUIRED",
            "error",
            "No complete scoring scale could be established.",
            f"{location}.levels",
            "Add explicit level values, complete overall thresholds, numeric level headers, or pass --allow-even-spacing.",
        )
        raw_weights = [_number(item.get("weight")) for item in (raw.get("criteria") or [])]
        if not raw_weights or any(value is None or value <= 0 for value in raw_weights):
            bag.add(
                "CRITERION_WEIGHT_REQUIRED",
                "error",
                "Criterion weights are missing, partial, or non-positive.",
                f"{location}.criteria",
                "Provide every weight; if all are absent, explicitly pass --allow-equal-weights.",
            )
        return None
    if (
        any(value < 0 or value > 1 for value in multipliers)
        or len(set(round(value, 12) for value in multipliers)) != len(multipliers)
    ):
        bag.add(
            "SCORING_METADATA_REQUIRED",
            "error",
            "Level multipliers must be unique values in [0, 1].",
            f"{location}.levels",
            "Correct the explicit scores or thresholds so the performance scale is unambiguous.",
        )
        return None
    if score_source == "approved_even_spacing":
        denominator = len(multipliers) - 1
        expected_even = [
            1.0 - index / denominator for index in range(len(multipliers))
        ]
        if any(
            abs(actual - expected) > 1e-9
            for actual, expected in zip(multipliers, expected_even)
        ):
            bag.add(
                "SCORING_METADATA_REQUIRED",
                "error",
                "Approved even-spacing provenance does not match an even scale.",
                f"{location}.levels",
                "Correct the scale or use explicit scoring provenance.",
            )
            return None
    if score_source == "approved_even_spacing" and not allow_even_spacing:
        bag.add(
            "SCORING_METADATA_REQUIRED",
            "error",
            "The contract declares approved even spacing without a matching approval record.",
            f"{location}.levels",
            "Set approvals.even_spacing=true or provide explicit scoring metadata.",
        )
        if any(
            item.get("weight_source") == "approved_equal_weights"
            for item in (raw.get("criteria") or [])
        ) and not allow_equal_weights:
            bag.add(
                "CRITERION_WEIGHT_REQUIRED",
                "error",
                "The contract declares approved equal weights without a current-run approval.",
                f"{location}.criteria",
                "Pass --allow-equal-weights for this run or provide explicit criterion weights.",
            )
        return None
    if score_source == "approved_even_spacing" and not any(
        item["code"] == "EVEN_SPACING_APPROVED" and item["location"] == f"{location}.levels"
        for item in bag.items
    ):
        bag.add(
            "EVEN_SPACING_APPROVED",
            "warning",
            "Even level spacing is present under an explicit approval record.",
            f"{location}.levels",
            "Review the generated multipliers and mapping before import.",
        )
    levels = [
        {
            "name": level_name,
            "sort_order": index,
            "multiplier": multiplier,
            "score_source": score_source,
            "extensions": {},
        }
        for index, (level_name, multiplier) in enumerate(
            zip(level_names, multipliers),
            start=1,
        )
    ]

    raw_criteria = raw.get("criteria") or []
    if not raw_criteria:
        bag.add(
            "SOURCE_SCHEMA_UNSUPPORTED",
            "error",
            "The rubric has no criteria.",
            f"{location}.criteria",
            "Provide at least one criterion row.",
        )
        return None
    criterion_names = [_clean_text(item.get("name", "")) for item in raw_criteria]
    duplicate_names = sorted(
        {
            value
            for value in criterion_names
            if value and sum(_identity_key(item) == _identity_key(value) for item in criterion_names) > 1
        }
    )
    if duplicate_names:
        bag.add(
            "DUPLICATE_CRITERION_NAME",
            "error",
            "The rubric contains duplicate criterion names.",
            f"{location}.criteria",
            "Give every criterion a unique name within the rubric.",
            duplicate_count=len(duplicate_names),
        )
        return None
    declared_weight_sources = [
        str(item.get("weight_source", ""))
        for item in raw_criteria
        if item.get("weight_source")
    ]
    declared_weight_source: str | None = None
    if adapter == AUTHORING_SCHEMA:
        unique_weight_sources = set(declared_weight_sources)
        if (
            len(declared_weight_sources) != len(raw_criteria)
            or len(unique_weight_sources) != 1
        ):
            bag.add(
                "CRITERION_WEIGHT_REQUIRED",
                "error",
                "Declared authoring weight provenance is incomplete or mixed.",
                f"{location}.criteria",
                "Use one uniform weight source for every criterion.",
            )
            return None
        declared_weight_source = next(iter(unique_weight_sources))
        if declared_weight_source == "extracted_cell_points":
            bag.add(
                "RUBRICS_ADAPTER_UNTRANSLATABLE",
                "error",
                "Declared authoring JSON cannot self-assert extracted weight provenance.",
                f"{location}.criteria",
                "Re-adapt the original coursecraft.rubrics/1 extraction contract.",
            )
            return None
    elif adapter == "coursecraft.rubrics/1":
        if (
            len(declared_weight_sources) != len(raw_criteria)
            or set(declared_weight_sources) != {"extracted_cell_points"}
        ):
            bag.add(
                "RUBRICS_ADAPTER_UNTRANSLATABLE",
                "error",
                "Extraction weight provenance is incomplete.",
                f"{location}.criteria",
                "Regenerate and re-adapt the extraction contract.",
            )
            return None
        declared_weight_source = "extracted_cell_points"
    elif declared_weight_sources:
        bag.add(
            "SOURCE_SCHEMA_UNSUPPORTED",
            "error",
            "This source adapter cannot declare weight provenance.",
            f"{location}.criteria",
            "Remove provenance fields and provide only explicit criterion weights.",
        )
        return None
    raw_weight_values = [item.get("weight") for item in raw_criteria]
    raw_weights = [_number(value) for value in raw_weight_values]
    weight_declared = [
        value is not None and str(value).strip() != ""
        for value in raw_weight_values
    ]
    weight_sources: list[str] = []
    if any(
        declared and (value is None or value <= 0)
        for declared, value in zip(weight_declared, raw_weights)
    ):
        bag.add(
            "CRITERION_WEIGHT_REQUIRED",
            "error",
            "A declared criterion weight is nonnumeric or non-positive.",
            f"{location}.criteria",
            "Correct every declared weight; invalid values cannot use the equal-weight fallback.",
        )
        return None
    if all(value is not None and value > 0 for value in raw_weights):
        weights = [float(value) for value in raw_weights if value is not None]
        weight_sources = [
            declared_weight_source or "explicit_weight"
        ] * len(raw_criteria)
    elif not any(weight_declared) and allow_equal_weights:
        base = round(100.0 / len(raw_criteria), 9)
        weights = [base for _ in raw_criteria]
        weights[-1] = round(100.0 - sum(weights[:-1]), 9)
        weight_sources = ["approved_equal_weights"] * len(raw_criteria)
        bag.add(
            "EQUAL_WEIGHTS_APPROVED",
            "warning",
            "Equal criterion weights were applied under explicit operator approval.",
            f"{location}.criteria",
            "Review the generated weights and mapping before import.",
        )
    else:
        bag.add(
            "CRITERION_WEIGHT_REQUIRED",
            "error",
            "Criterion weights are missing, partial, or non-positive.",
            f"{location}.criteria",
            "Provide every weight; if all are absent, explicitly pass --allow-equal-weights.",
        )
        return None
    if "approved_equal_weights" in weight_sources and not allow_equal_weights:
        bag.add(
            "CRITERION_WEIGHT_REQUIRED",
            "error",
            "The contract declares approved equal weights without a matching approval record.",
            f"{location}.criteria",
            "Set approvals.equal_weights=true or provide explicit criterion weights.",
        )
        return None
    if "approved_equal_weights" in weight_sources and any(
        abs(value - weights[0]) > 1e-9 for value in weights[1:]
    ):
        bag.add(
            "CRITERION_WEIGHT_REQUIRED",
            "error",
            "Approved equal-weight provenance does not match equal weights.",
            f"{location}.criteria",
            "Correct the weights or use explicit weight provenance.",
        )
        return None
    if "approved_equal_weights" in weight_sources and not any(
        item["code"] == "EQUAL_WEIGHTS_APPROVED" and item["location"] == f"{location}.criteria"
        for item in bag.items
    ):
        bag.add(
            "EQUAL_WEIGHTS_APPROVED",
            "warning",
            "Equal criterion weights are present under an explicit approval record.",
            f"{location}.criteria",
            "Review the generated weights and mapping before import.",
        )
    if abs(sum(weights) - 100.0) > 1e-6:
        bag.add(
            "CRITERION_WEIGHT_REQUIRED",
            "error",
            f"Criterion weights total {sum(weights):g}, not 100.",
            f"{location}.criteria",
            "Correct explicit criterion weights so they total exactly 100.",
        )
        return None
    criteria: list[dict[str, Any]] = []
    for index, (criterion, criterion_name, weight, weight_source) in enumerate(
        zip(raw_criteria, criterion_names, weights, weight_sources),
        start=1,
    ):
        criterion_location = criterion.get("_location", f"{location}.criterion[{index}]")
        if not criterion_name:
            bag.add(
                "SOURCE_SCHEMA_UNSUPPORTED",
                "error",
                "A criterion has no name.",
                f"{criterion_location}.name",
                "Provide a non-empty criterion name.",
            )
            continue
        descriptions = criterion.get("descriptions", criterion.get("levels", {}))
        if not isinstance(descriptions, dict):
            descriptions = {}
        normalized_descriptions = {
            level_name: _paragraph_html(descriptions.get(level_name, ""))
            for level_name in level_names
        }
        missing = [name for name, value in normalized_descriptions.items() if not value]
        if missing:
            bag.add(
                "SOURCE_SCHEMA_UNSUPPORTED",
                "error",
                "A criterion is missing one or more level descriptions.",
                f"{criterion_location}.descriptions",
                "Provide exactly one non-empty authored description for every performance level.",
                missing_count=len(missing),
            )
            continue
        criteria.append(
            {
                "name": criterion_name,
                "sort_order": index,
                "weight": weight,
                "weight_source": weight_source,
                "descriptions": normalized_descriptions,
                "extensions": {},
            }
        )
    if len(criteria) != len(raw_criteria):
        return None

    overall_map = _metadata_pairs(raw.get("overall_levels"))
    if overall_map is not None and set(overall_map) == set(level_names):
        overall_source = raw.get("overall_source") or "explicit_overall_threshold"
        thresholds = overall_map
    else:
        overall_source = "derived_from_level_multiplier"
        thresholds = {
            level["name"]: round(float(level["multiplier"]) * 100.0, 9)
            for level in levels
        }
    if (
        set(thresholds) != set(level_names)
        or any(value < 0 or value > 100 for value in thresholds.values())
        or min(thresholds.values()) != 0
        or len(set(round(value, 9) for value in thresholds.values())) != len(thresholds)
    ):
        bag.add(
            "SCORING_METADATA_REQUIRED",
            "error",
            "Overall thresholds must uniquely cover all levels, stay in [0, 100], and begin at 0.",
            f"{location}.overall_levels",
            "Correct the complete overall threshold metadata.",
        )
        return None
    overall_levels = [
        {
            "name": level["name"],
            "sort_order": level["sort_order"],
            "range_start_value": thresholds[level["name"]],
            "source": overall_source,
            "extensions": {},
        }
        for level in sorted(levels, key=lambda item: thresholds[item["name"]])
    ]
    return {
        "name": name,
        "kind": _clean_text(raw.get("kind", "other")) or "other",
        "source_reference": (
            _clean_text(raw.get("source_reference"))
            if raw.get("source_reference") is not None
            else None
        ),
        "levels": levels,
        "criteria": criteria,
        "overall_levels": overall_levels,
        "extensions": {},
    }


def normalize_source(
    input_path: Path,
    *,
    allow_even_spacing: bool = False,
    allow_equal_weights: bool = False,
    manifest_identifier: str | None = None,
    resource_prefix: str | None = None,
    source_label: str | None = None,
) -> dict[str, Any]:
    """Adapt and normalize one supported source, or raise AuthoringRefusal."""

    bag = DiagnosticBag([])
    raw_path = input_path.expanduser()
    if raw_path.is_symlink():
        bag.add(
            "SOURCE_PATH_UNSAFE",
            "error",
            "Symlink rubric sources are not accepted.",
            "source",
            "Stage the source as a regular file and retry.",
        )
        raise AuthoringRefusal(bag.items)
    if not raw_path.exists():
        bag.add(
            "SOURCE_PATH_MISSING",
            "error",
            "The rubric source does not exist.",
            "source",
            "Stage a readable regular source file and retry.",
        )
        raise AuthoringRefusal(bag.items)
    if not raw_path.is_file():
        bag.add(
            "SOURCE_PATH_UNSAFE",
            "error",
            "The rubric source is not a regular file.",
            "source",
            "Stage a readable regular source file and retry.",
        )
        raise AuthoringRefusal(bag.items)
    try:
        with raw_path.open("rb") as source_handle:
            source_handle.read(1)
    except OSError:
        bag.add(
            "SOURCE_UNREADABLE",
            "error",
            "The rubric source is not readable.",
            "source",
            "Correct local file permissions and retry.",
        )
        raise AuthoringRefusal(bag.items)
    path = raw_path.resolve()
    suffix = path.suffix.lower()
    source_bytes = path.stat().st_size
    source_limit = (
        MAX_DOCX_SOURCE_BYTES if suffix == ".docx" else MAX_TEXT_SOURCE_BYTES
    )
    if source_bytes > source_limit:
        bag.add(
            "SOURCE_SIZE_LIMIT",
            "error",
            "The rubric source exceeds the producer size limit.",
            "source",
            "Reduce the source size and retry.",
            bytes=source_bytes,
            limit=source_limit,
        )
        raise AuthoringRefusal(bag.items)
    source_data: dict[str, Any] = {}
    if suffix in {".md", ".markdown", ".txt"}:
        try:
            raw_rubrics = _parse_markdown(path, bag)
        except (OSError, UnicodeError) as exc:
            bag.add(
                "SOURCE_SCHEMA_UNSUPPORTED",
                "error",
                "The Markdown source is not readable UTF-8 text.",
                "source",
                "Save the source as UTF-8 text and retry.",
                error_type=type(exc).__name__,
            )
            raw_rubrics = []
        adapter = "markdown_table"
        media_kind = "markdown"
    elif suffix == ".docx":
        if not _preflight_docx_archive(path, bag):
            raw_rubrics = []
        else:
            try:
                raw_rubrics = _parse_docx(path, bag)
            except Exception:
                bag.add(
                    "SOURCE_SCHEMA_UNSUPPORTED",
                    "error",
                    "The Word source could not be inspected.",
                    "source",
                    "Provide an intact DOCX containing straightforward rubric tables.",
                )
                raw_rubrics = []
        adapter = "docx_table"
        media_kind = "docx"
    elif suffix == ".json":
        raw_rubrics, adapter, source_data = _parse_json(path, bag)
        media_kind = "json"
    else:
        bag.add(
            "SOURCE_SCHEMA_UNSUPPORTED",
            "error",
            f"Unsupported input extension: {suffix or '(none)'}.",
            "source",
            "Use DOCX, Markdown, text, or JSON.",
        )
        raw_rubrics = []
        adapter = "legacy_builder_json"
        media_kind = "json"

    names = [_clean_text(rubric.get("name", "")) for rubric in raw_rubrics]
    duplicates = sorted(
        {
            value
            for value in names
            if value and sum(_identity_key(item) == _identity_key(value) for item in names) > 1
        }
    )
    if duplicates:
        bag.add(
            "DUPLICATE_RUBRIC_NAME",
            "error",
            "The source contains duplicate rubric names.",
            "rubrics",
            "Give every rubric a unique name.",
            duplicate_count=len(duplicates),
        )
    ids = [
        str(rubric.get(key))
        for rubric in raw_rubrics
        for key in ("source_id", "resource_code")
        if rubric.get(key)
    ]
    collisions = sorted({value for value in ids if ids.count(value) > 1})
    if collisions:
        bag.add(
            "IDENTIFIER_COLLISION",
            "error",
            "The source contains colliding explicit identifiers.",
            "rubrics",
            "Remove duplicate source/resource identifiers before packaging.",
            collision_count=len(collisions),
        )

    source_approvals = (
        source_data.get("approvals", {})
        if adapter == AUTHORING_SCHEMA and isinstance(source_data, dict)
        else {}
    )
    if source_approvals:
        bag.add(
            "PRIOR_APPROVAL_OBSERVED",
            "info",
            "The input contract contains prior approval facts; they do not authorize this run.",
            "source.approvals",
            "Supply the current CLI approval flags if the prior contract uses an approved fallback.",
            prior_approvals={
                "even_spacing": bool(source_approvals.get("even_spacing", False)),
                "equal_weights": bool(source_approvals.get("equal_weights", False)),
            },
        )
    effective_even_spacing = bool(allow_even_spacing)
    effective_equal_weights = bool(allow_equal_weights)
    normalized_rubrics = [
        normalized
        for raw in raw_rubrics
        if (
            normalized := _normalize_rubric(
                raw,
                bag,
                allow_even_spacing=effective_even_spacing,
                allow_equal_weights=effective_equal_weights,
                adapter=adapter,
            )
        )
        is not None
    ]
    package_source = source_data.get("package", {}) if isinstance(source_data, dict) else {}
    try:
        rubric_id_start = int(package_source.get("rubric_id_start", 1))
        resource_code_start = int(package_source.get("resource_code_start", 980001))
        level_id_base_start = int(package_source.get("level_id_base_start", 660000))
    except (TypeError, ValueError):
        bag.add(
            "IDENTIFIER_COLLISION",
            "error",
            "Package allocation seeds must be integers.",
            "package",
            "Provide positive integer rubric_id_start, resource_code_start, and level_id_base_start values.",
        )
        rubric_id_start, resource_code_start, level_id_base_start = 1, 980001, 660000
    package = {
        "manifest_identifier": _safe_token(
            manifest_identifier
            or package_source.get("manifest_identifier", "D2L_RUBRICS_PACKAGE"),
            "RUBRICS_PACKAGE",
        ),
        "resource_prefix": _safe_token(
            resource_prefix
            or package_source.get("resource_prefix", "COURSECRAFT_RUBRIC"),
            "RUBRIC",
        ),
        "rubric_id_start": rubric_id_start,
        "resource_code_start": resource_code_start,
        "level_id_base_start": level_id_base_start,
    }
    if isinstance(package_source.get("orgunit"), dict):
        package["orgunit"] = {
            str(key): str(value)
            for key, value in package_source["orgunit"].items()
            if value is not None
        }
    allocated_level_ids = [
        level_id_base_start + ((rubric_index + 1) * 100) + level_index
        for rubric_index, rubric in enumerate(normalized_rubrics)
        for level_index, _ in enumerate(rubric["levels"], start=1)
    ]
    if len(allocated_level_ids) != len(set(allocated_level_ids)):
        bag.add(
            "IDENTIFIER_COLLISION",
            "error",
            "The configured level-ID allocation collides across rubrics.",
            "package.level_id_base_start",
            "Use fewer than 101 levels per rubric or revise the canonical allocation strategy.",
        )
    prior_lineage: list[dict[str, Any]] = []
    if adapter == AUTHORING_SCHEMA and isinstance(source_data.get("source"), dict):
        existing_lineage = source_data.get("extensions", {}).get("source_lineage", [])
        if isinstance(existing_lineage, list):
            prior_lineage.extend(
                item for item in existing_lineage if isinstance(item, dict)
            )
        prior_lineage.append(dict(source_data["source"]))
    contract_extensions: dict[str, Any] = {}
    if source_approvals:
        contract_extensions["prior_approvals"] = {
            "even_spacing": bool(source_approvals.get("even_spacing", False)),
            "equal_weights": bool(source_approvals.get("equal_weights", False)),
        }
    if prior_lineage:
        contract_extensions["source_lineage"] = prior_lineage
    contract = {
        "schema": AUTHORING_SCHEMA,
        "package": package,
        "source": {
            "media_kind": media_kind,
            "label": _generic_source_label(source_label),
            "sha256": _sha256_file(path),
            "adapter": adapter,
            "extensions": {"bytes": source_bytes},
        },
        "approvals": {
            "even_spacing": effective_even_spacing,
            "equal_weights": effective_equal_weights,
        },
        "rubrics": normalized_rubrics,
        "diagnostics": bag.items,
        "notes": [
            "This contract produces rubric objects only.",
            "Activity attachment remains manual and explicit.",
        ],
        "extensions": contract_extensions,
    }
    if not normalized_rubrics and not bag.has_errors:
        bag.add(
            "SOURCE_SCHEMA_UNSUPPORTED",
            "error",
            "No runnable rubric remained after normalization.",
            "rubrics",
            "Correct the source and run preflight again.",
        )
    if not bag.has_errors:
        schema = json.loads(SCHEMA_PATH.read_text(encoding="utf-8"))
        schema_errors = sorted(Draft7Validator(schema).iter_errors(contract), key=lambda item: list(item.path))
        for error in schema_errors:
            location = ".".join(str(item) for item in error.path) or "contract"
            bag.add(
                "SOURCE_SCHEMA_UNSUPPORTED",
                "error",
                "The normalized contract does not satisfy the authoring schema.",
                location,
                "Correct the source so it satisfies coursecraft.rubric_authoring/1.",
                validator=error.validator,
            )
    if bag.has_errors:
        raise AuthoringRefusal(bag.items)
    return contract


def preflight_summary(contract: dict[str, Any]) -> dict[str, Any]:
    return {
        "schema": "coursecraft.rubric_authoring_preflight/1",
        "status": "ok",
        "source": contract["source"],
        "approvals": contract["approvals"],
        "rubric_count": len(contract["rubrics"]),
        "rubrics": [
            {
                "name": rubric["name"],
                "levels": [
                    {
                        "name": level["name"],
                        "multiplier": level["multiplier"],
                        "score_source": level["score_source"],
                        "overall_threshold": next(
                            item["range_start_value"]
                            for item in rubric["overall_levels"]
                            if item["name"] == level["name"]
                        ),
                        "overall_threshold_source": next(
                            item["source"]
                            for item in rubric["overall_levels"]
                            if item["name"] == level["name"]
                        ),
                    }
                    for level in rubric["levels"]
                ],
                "criteria": [
                    {
                        "name": criterion["name"],
                        "weight": criterion["weight"],
                        "weight_source": criterion["weight_source"],
                    }
                    for criterion in rubric["criteria"]
                ],
            }
            for rubric in contract["rubrics"]
        ],
        "diagnostics": contract["diagnostics"],
    }


def _mapping_markdown(contract: dict[str, Any]) -> str:
    def prose(value: Any) -> str:
        text = _clean_text(value)
        for char in ("\\", "`", "*", "_", "[", "]", "<", ">", "#"):
            text = text.replace(char, f"\\{char}")
        return text

    def table(value: Any) -> str:
        return prose(value).replace("|", r"\|")

    lines = [
        "# Rubric Weave Mapping and Review",
        "",
        f"- Source: {prose(contract['source']['label'])}",
        f"- Adapter: `{contract['source']['adapter']}`",
        f"- Source SHA-256: `{contract['source']['sha256']}`",
        f"- Even spacing approved: `{str(contract['approvals']['even_spacing']).lower()}`",
        f"- Equal weights approved: `{str(contract['approvals']['equal_weights']).lower()}`",
        "- Package scope: rubric objects only; no activity attachment payloads.",
        "",
    ]
    for rubric in contract["rubrics"]:
        lines.extend(
            [
                f"## {prose(rubric['name'])}",
                "",
                f"- Source reference: {prose(rubric.get('source_reference') or '(none)')}",
                "",
                "| Level | Multiplier | Score source | Overall threshold |",
                "| --- | ---: | --- | ---: |",
            ]
        )
        thresholds = {
            item["name"]: item["range_start_value"]
            for item in rubric["overall_levels"]
        }
        for level in rubric["levels"]:
            lines.append(
                f"| {table(level['name'])} | {level['multiplier']:g} | {table(level['score_source'])} "
                f"| {thresholds[level['name']]:g} |"
            )
        lines.extend(
            [
                "",
                "| Criterion | Weight | Weight source |",
                "| --- | ---: | --- |",
            ]
        )
        for criterion in rubric["criteria"]:
            lines.append(
                f"| {table(criterion['name'])} | {criterion['weight']:g} | "
                f"{table(criterion['weight_source'])} |"
            )
        lines.append("")
    if contract["diagnostics"]:
        lines.extend(["## Diagnostics", ""])
        for item in contract["diagnostics"]:
            lines.append(
                f"- `{item['code']}` ({item['severity']}): {item['message']} "
                f"Location: `{item['location']}`"
            )
        lines.append("")
    return "\n".join(lines)


def _semantic_roundtrip(contract: dict[str, Any], xml_path: Path) -> list[str]:
    """Compare authored semantics through the canonical Unravel extractor."""

    errors: list[str] = []
    extracted = rubrics_to_records(xml_path)
    if extracted.get("diagnostics"):
        errors.extend(
            f"Canonical extractor returned diagnostic {index}."
            for index, _message in enumerate(extracted["diagnostics"], start=1)
        )
    extracted_rubrics = extracted.get("rubrics", [])
    if len(extracted_rubrics) != len(contract["rubrics"]):
        return ["Rubric count changed during XML generation."]
    package = contract["package"]
    for rubric_offset, (expected, actual) in enumerate(
        zip(contract["rubrics"], extracted_rubrics)
    ):
        if actual["name"] != expected["name"]:
            errors.append(f"Rubric {rubric_offset + 1} name mismatch.")
        if actual["id"] != str(int(package["rubric_id_start"]) + rubric_offset):
            errors.append(f"Rubric {rubric_offset + 1} ID mismatch.")
        expected_resource_code = (
            f"{package['resource_prefix']}-"
            f"{int(package['resource_code_start']) + rubric_offset}"
        )
        if actual["resource_code"] != expected_resource_code:
            errors.append(f"Rubric {rubric_offset + 1} resource code mismatch.")
        actual_level_names = [item["name"] for item in actual["levels"]]
        expected_level_names = [item["name"] for item in expected["levels"]]
        if actual_level_names != expected_level_names:
            errors.append(f"Rubric {rubric_offset + 1} level order mismatch.")
            continue
        if [item["sort_order"] for item in actual["levels"]] != [
            item["sort_order"] for item in expected["levels"]
        ]:
            errors.append(f"Rubric {rubric_offset + 1} level sort-order mismatch.")
        actual_criteria = actual["criteria"]
        if [item["name"] for item in actual_criteria] != [
            item["name"] for item in expected["criteria"]
        ]:
            errors.append(f"Rubric {rubric_offset + 1} criterion order mismatch.")
            continue
        multipliers = {
            level["name"]: float(level["multiplier"])
            for level in expected["levels"]
        }
        for criterion_offset, (expected_criterion, actual_criterion) in enumerate(
            zip(expected["criteria"], actual_criteria)
        ):
            by_name = {
                cell["level_name"]: cell
                for cell in actual_criterion["cells"]
            }
            for level_offset, level_name in enumerate(expected_level_names):
                cell = by_name.get(level_name)
                if cell is None:
                    errors.append(
                        f"Rubric {rubric_offset + 1} criterion {criterion_offset + 1} "
                        f"lost level {level_offset + 1}."
                    )
                    continue
                actual_score = cell["points"]
                expected_score = expected_criterion["weight"] * multipliers[level_name]
                if actual_score is None or abs(actual_score - expected_score) > 1e-7:
                    errors.append(
                        f"Rubric {rubric_offset + 1} criterion {criterion_offset + 1} "
                        f"cell {level_offset + 1} score mismatch."
                    )
                expected_text = clean_xml_text(
                    expected_criterion["descriptions"][level_name]
                )
                if cell["description"] != expected_text:
                    errors.append(
                        f"Rubric {rubric_offset + 1} criterion {criterion_offset + 1} "
                        f"cell {level_offset + 1} description mismatch."
                    )
        actual_thresholds = {
            item["name"]: item["range_start_value"]
            for item in actual["overall_levels"]
        }
        expected_thresholds = {
            item["name"]: float(item["range_start_value"])
            for item in expected["overall_levels"]
        }
        if set(actual_thresholds) != set(expected_thresholds) or any(
            actual_thresholds[name] is None
            or abs(float(actual_thresholds[name]) - expected_thresholds[name]) > 1e-7
            for name in expected_thresholds
            if name in actual_thresholds
        ):
            errors.append(f"Rubric {rubric_offset + 1} overall threshold mismatch.")
    return errors


def _artifact(path: Path, root: Path, media_type: str, contract: str | None, role: str) -> dict[str, Any]:
    return {
        "path": path.relative_to(root).as_posix(),
        "bytes": path.stat().st_size,
        "sha256": _sha256_file(path),
        "media_type": media_type,
        "contract": contract,
        "extensions": {"role": role},
    }


def _receipt_diagnostic(item: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": item["id"],
        "code": item["code"],
        "severity": item["severity"],
        "location": item["location"],
        "extensions": (
            {"approval_recorded": True}
            if item["code"] in {"EVEN_SPACING_APPROVED", "EQUAL_WEIGHTS_APPROVED"}
            else {}
        ),
    }


def assert_safe_output_target(
    output_dir: Path,
    *,
    input_path: Path,
    context_dir: Path | None = None,
) -> Path:
    """Resolve and validate an authoring output target before replacement."""

    sources = [input_path]
    if context_dir is not None:
        sources.append(context_dir)
    return assert_safe_replace_target(
        output_dir,
        protected_sources=tuple(sources),
    )


def _producer_identity() -> dict[str, Any]:
    """Return content-minimized local git identity without paths or remote URLs."""

    repo_root = Path(__file__).resolve().parents[1]

    def git(*args: str, check: bool = True) -> str:
        result = subprocess.run(
            ["git", *args],
            cwd=repo_root,
            text=True,
            capture_output=True,
            check=check,
        )
        return result.stdout.strip()

    code_digests = {
        path.name: _sha256_file(path)
        for path in (
            Path(__file__).resolve(),
            Path(__file__).resolve().with_name("rubric_package_lib.py"),
            Path(__file__).resolve().with_name("extract_rubrics_to_workbook.py"),
        )
    }
    schema_digests = {
        path.name: _sha256_file(path)
        for path in (SCHEMA_PATH, RUN_SCHEMA_PATH, RUBRICS_SCHEMA_PATH)
    }
    try:
        commit = git("rev-parse", "HEAD")
        ref = git(
            "symbolic-ref",
            "--quiet",
            "--short",
            "HEAD",
            check=False,
        ) or None
        dirty = bool(git("status", "--porcelain", "--untracked-files=normal"))
    except (OSError, subprocess.CalledProcessError):
        return {
            "component": "coursecraft-workbench-rubric-authoring",
            "identity_state": "unknown",
            "version": None,
            "repository": "coursecraft-workbench",
            "ref": None,
            "commit": None,
            "dirty": None,
            "extensions": {
                "reason": "git_identity_unavailable",
                "code_digests": code_digests,
                "schema_digests": schema_digests,
            },
        }
    return {
        "component": "coursecraft-workbench-rubric-authoring",
        "identity_state": "git",
        "version": None,
        "repository": "coursecraft-workbench",
        "ref": ref,
        "commit": commit,
        "dirty": dirty,
        "extensions": {
            "code_digests": code_digests,
            "schema_digests": schema_digests,
        },
    }


def _build_weave_outputs_direct(
    input_path: Path,
    output_dir: Path,
    *,
    allow_even_spacing: bool = False,
    allow_equal_weights: bool = False,
    context_dir: Path | None = None,
    cli_overrides: dict[str, str] | None = None,
    manifest_identifier: str | None = None,
    resource_prefix: str | None = None,
    source_label: str | None = None,
    force: bool = False,
) -> dict[str, Path]:
    contract = normalize_source(
        input_path,
        allow_even_spacing=allow_even_spacing,
        allow_equal_weights=allow_equal_weights,
        manifest_identifier=manifest_identifier,
        resource_prefix=resource_prefix,
        source_label=source_label,
    )
    output_dir = assert_safe_output_target(
        output_dir,
        input_path=input_path,
        context_dir=context_dir,
    )
    if output_dir.exists():
        if not force:
            raise ValueError(f"Output directory already exists: {output_dir}. Re-run with --force to replace it.")
        shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True)
    context_defaults = {
        "identifier": "COURSECRAFT_RUBRICS",
        "default_nav": "Course Default",
        "default_homepage": "Course Home",
        "title": "Rubric Package",
        "keyword": "Rubric Package",
    }
    context_defaults.update(
        {
            key: str(value)
            for key, value in contract["package"].get("orgunit", {}).items()
            if key in context_defaults and value
        }
    )
    context_defaults.update({key: value for key, value in (cli_overrides or {}).items() if value})
    context = merge_context(
        contract,
        context_dir=context_dir,
        cli_overrides=context_defaults if context_dir is None else cli_overrides,
    )
    if context["identifier"] == "res_rubrics":
        raise AuthoringRefusal(
            [
                {
                    "id": "diag-0001",
                    "code": "IDENTIFIER_COLLISION",
                    "severity": "error",
                    "message": "The org-unit resource identifier collides with the rubric resource identifier.",
                    "location": "context.identifier",
                    "remediation": "Choose an org-unit identifier other than 'res_rubrics'.",
                    "extensions": {},
                }
            ]
        )
    package_dir = output_dir / "package"
    xml_path = package_dir / "rubrics_d2l.xml"
    write_xml(xml_path, build_rubrics_xml(contract))
    write_xml(package_dir / "imsmanifest.xml", build_manifest_xml(context, contract))
    write_xml(package_dir / "orgunitconfig" / "orgunitconfig.xml", build_orgunit_xml(context))
    zip_path = output_dir / "rubric_package.zip"
    zip_package(package_dir, zip_path)
    top_xml = output_dir / "rubrics_d2l.xml"
    shutil.copyfile(xml_path, top_xml)
    normalized_path = output_dir / "normalized_rubric_authoring.json"
    normalized_path.write_bytes(_json_bytes(contract))
    mapping_path = output_dir / "rubric_mapping.md"
    mapping_path.write_text(_mapping_markdown(contract), encoding="utf-8")

    folder_errors, folder_warnings, folder_summary = validate_package_path(package_dir)
    zip_errors, zip_warnings, zip_summary = validate_package_path(zip_path)
    errors = folder_errors + zip_errors
    warnings = folder_warnings + zip_warnings
    summary = zip_summary or folder_summary
    roundtrip_errors = _semantic_roundtrip(contract, xml_path)
    diagnostics = list(contract["diagnostics"])
    if errors or roundtrip_errors:
        for message in errors + roundtrip_errors:
            diagnostics.append(
                {
                    "id": f"diag-{len(diagnostics) + 1:04d}",
                    "code": "SEMANTIC_ROUNDTRIP_MISMATCH",
                    "severity": "error",
                    "message": message,
                    "location": "outputs.rubric_package",
                    "remediation": "Do not import this package; inspect the canonical producer.",
                    "extensions": {},
                }
            )
        raise AuthoringRefusal(diagnostics, "Generated package failed validation.")
    diagnostics.append(
        {
            "id": f"diag-{len(diagnostics) + 1:04d}",
            "code": "SEMANTIC_ROUNDTRIP_OK",
            "severity": "info",
            "message": "Package conformance and semantic round-trip validation passed.",
            "location": "outputs.rubric_package",
            "remediation": "No action required.",
            "extensions": {"rubric_count": summary.get("rubric_count"), "warnings": warnings},
        }
    )
    diagnostics_path = output_dir / "diagnostics.json"
    diagnostics_path.write_bytes(
        _json_bytes(
            {
                "schema": "coursecraft.diagnostics/1",
                "status": "ok",
                "diagnostics": diagnostics,
            }
        )
    )
    conversion_review: Path | None = None
    if contract["source"]["media_kind"] == "docx":
        conversion_review = output_dir / "conversion_review.md"
        conversion_review.write_text(
            "# DOCX Conversion Review\n\n"
            "The strict DOCX adapter accepted only rectangular, unmerged, non-nested tables. "
            "The normalized names, levels, scores, and weights are repeated below for operator review. "
            "Read `normalized_rubric_authoring.json` for every preserved description.\n\n"
            + _mapping_markdown(contract),
            encoding="utf-8",
        )

    artifacts = [
        _artifact(zip_path, output_dir, "application/zip", None, "rubric_import_package"),
        _artifact(top_xml, output_dir, "application/xml", None, "rubrics_xml_companion"),
        _artifact(
            normalized_path,
            output_dir,
            "application/json",
            AUTHORING_SCHEMA,
            "normalized_authoring_contract",
        ),
        _artifact(mapping_path, output_dir, "text/markdown", None, "mapping_review"),
        _artifact(
            diagnostics_path,
            output_dir,
            "application/json",
            "coursecraft.diagnostics/1",
            "diagnostics",
        ),
    ]
    if conversion_review:
        artifacts.append(
            _artifact(conversion_review, output_dir, "text/markdown", None, "docx_conversion_review")
        )
    source_size = int(contract["source"]["extensions"]["bytes"])
    source_identity_bytes = bytes.fromhex(contract["source"]["sha256"])
    producer_identity = _producer_identity()
    options_identity = {
        "allow_even_spacing": allow_even_spacing,
        "allow_equal_weights": allow_equal_weights,
        "source_label": contract["source"]["label"],
        "package": contract["package"],
        "context": context,
        "producer": producer_identity,
    }
    run_digest = _sha256_bytes(source_identity_bytes + _json_bytes(options_identity))
    authoring_schema_sha = _sha256_file(SCHEMA_PATH)
    run_schema_sha = _sha256_file(RUN_SCHEMA_PATH)
    minimized_parameters = {
        "allow_even_spacing": bool(allow_even_spacing),
        "allow_equal_weights": bool(allow_equal_weights),
        "source_label_sha256": _sha256_bytes(
            contract["source"]["label"].encode("utf-8")
        ),
        "package_sha256": _sha256_bytes(_json_bytes(contract["package"])),
        "context_sha256": _sha256_bytes(_json_bytes(context)),
        "identity_parameters_sha256": _sha256_bytes(_json_bytes(options_identity)),
    }
    contract_receipts = [
        {
            "schema": AUTHORING_SCHEMA,
            "schema_path": "workspace/reference/schemas/rubrics/rubric_authoring_schema.json",
            "sha256": authoring_schema_sha,
            "extensions": {"role": "normalized_output"},
        },
        {
            "schema": RUN_SCHEMA,
            "schema_path": "workspace/reference/schemas/course/run_identity_schema.json",
            "sha256": run_schema_sha,
            "extensions": {"role": "run_receipt"},
        },
    ]
    if contract["source"]["adapter"] == "coursecraft.rubrics/1":
        contract_receipts.append(
            {
                "schema": "coursecraft.rubrics/1",
                "schema_path": "workspace/reference/schemas/rubrics/rubrics_schema.json",
                "sha256": _sha256_file(RUBRICS_SCHEMA_PATH),
                "extensions": {
                    "role": "input_contract",
                    "input_document_sha256": contract["source"]["sha256"],
                },
            }
        )
    lineage_receipt = [
        {
            "sha256": item.get("sha256"),
            "media_kind": item.get("media_kind"),
            "adapter": item.get("adapter"),
        }
        for item in contract.get("extensions", {}).get("source_lineage", [])
        if isinstance(item, dict) and item.get("sha256")
    ]
    receipt_path = output_dir / "run_receipt.json"
    receipt = {
        "schema": RUN_SCHEMA,
        "run_id": f"cc:run:{run_digest[:24]}",
        "status": "ok",
        "started_at": None,
        "finished_at": None,
        "source": {
            "source_lineage_key": f"cc:lineage:{contract['source']['sha256']}",
            "source_instance_key": f"cc:source:{contract['source']['sha256']}",
            "lineage_state": "resolved",
            "lineage_basis": "source-file-sha256",
            "source_name": "rubric-source",
            "logical_fingerprint": {
                "algorithm": "sha256",
                "digest": contract["source"]["sha256"],
                "scope": "source_file",
                "file_count": 1,
                "bytes": source_size,
                "extensions": {},
            },
            "transport_fingerprint": {
                "algorithm": "sha256",
                "digest": contract["source"]["sha256"],
                "scope": "source_file_bytes",
                "file_count": 1,
                "bytes": source_size,
                "extensions": {},
            },
            "observed_identity": {
                "media_kind": contract["source"]["media_kind"],
                "adapter": contract["source"]["adapter"],
                "source_label_sha256": minimized_parameters["source_label_sha256"],
            },
            "extensions": {"source_lineage": lineage_receipt},
        },
        "producer": producer_identity,
        "contracts": contract_receipts,
        "parameters": minimized_parameters,
        "steps": [
            {
                "name": "normalize_source",
                "status": "completed",
                "started_at": None,
                "finished_at": None,
                "artifact_paths": ["normalized_rubric_authoring.json", "diagnostics.json"],
                "diagnostic_ids": [item["id"] for item in diagnostics],
                "notes": [],
                "extensions": {},
            },
            {
                "name": "build_rubric_only_package",
                "status": "completed",
                "started_at": None,
                "finished_at": None,
                "artifact_paths": ["rubric_package.zip", "rubrics_d2l.xml"],
                "diagnostic_ids": [],
                "notes": ["No activity payloads were generated or modified."],
                "extensions": {},
            },
            {
                "name": "validate_and_roundtrip",
                "status": "completed",
                "started_at": None,
                "finished_at": None,
                "artifact_paths": ["rubric_package.zip", "rubric_mapping.md"],
                "diagnostic_ids": [diagnostics[-1]["id"]],
                "notes": [],
                "extensions": {},
            },
        ],
        "emitted_files": artifacts,
        "receipt_path": "run_receipt.json",
        "diagnostics": [_receipt_diagnostic(item) for item in diagnostics],
        "extensions": {
            "approvals": contract["approvals"],
            "package_members": list(EXPECTED_PACKAGE_MEMBERS),
        },
    }
    run_schema = json.loads(RUN_SCHEMA_PATH.read_text(encoding="utf-8"))
    receipt_errors = list(Draft7Validator(run_schema).iter_errors(receipt))
    if receipt_errors:
        raise ValueError(f"Internal run receipt validation failed: {receipt_errors[0].message}")
    receipt_path.write_bytes(_json_bytes(receipt))
    result = {
        "output_dir": output_dir,
        "package_dir": package_dir,
        "zip_path": zip_path,
        "rubrics_xml_path": top_xml,
        "normalized_json_path": normalized_path,
        "mapping_path": mapping_path,
        "diagnostics_path": diagnostics_path,
        "receipt_path": receipt_path,
    }
    if conversion_review:
        result["conversion_review_path"] = conversion_review
    return result


def build_weave_outputs(
    input_path: Path,
    output_dir: Path,
    *,
    allow_even_spacing: bool = False,
    allow_equal_weights: bool = False,
    context_dir: Path | None = None,
    cli_overrides: dict[str, str] | None = None,
    manifest_identifier: str | None = None,
    resource_prefix: str | None = None,
    source_label: str | None = None,
    force: bool = False,
) -> dict[str, Path]:
    """Build transactionally, preserving any prior output until success."""

    final_root = assert_safe_output_target(
        output_dir,
        input_path=input_path,
        context_dir=context_dir,
    )
    if final_root.exists() and not force:
        raise ValueError("The output directory already exists; use --force to replace it.")
    final_root.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(
        prefix=f".{final_root.name}.candidate-",
        dir=final_root.parent,
    ) as temporary:
        candidate_root = Path(temporary)
        candidate = _build_weave_outputs_direct(
            input_path,
            candidate_root,
            allow_even_spacing=allow_even_spacing,
            allow_equal_weights=allow_equal_weights,
            context_dir=context_dir,
            cli_overrides=cli_overrides,
            manifest_identifier=manifest_identifier,
            resource_prefix=resource_prefix,
            source_label=source_label,
            force=True,
        )
        backup: Path | None = None
        try:
            if final_root.exists():
                backup = final_root.parent / f".{final_root.name}.previous"
                if backup.exists():
                    raise ValueError("A recovery backup already exists; refusing replacement.")
                final_root.rename(backup)
            candidate_root.rename(final_root)
        except Exception:
            if backup is not None and backup.exists() and not final_root.exists():
                backup.rename(final_root)
            raise
        if backup is not None and backup.exists():
            # The new output is already atomically in place. Backup cleanup is
            # best-effort so a permissions race cannot turn a completed swap
            # into a reported failed run.
            shutil.rmtree(backup, ignore_errors=True)
        return {
            key: (
                final_root
                if key == "output_dir"
                else final_root / path.relative_to(candidate_root)
            )
            for key, path in candidate.items()
        }

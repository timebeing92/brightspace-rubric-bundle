from __future__ import annotations

import json
import hashlib
import subprocess
import sys
import tempfile
import xml.etree.ElementTree as ET
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import pytest
from docx import Document
from jsonschema import Draft7Validator


REPO_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = REPO_ROOT / "scripts"
FIXTURES = REPO_ROOT / "tests" / "fixtures" / "rubric_authoring"
sys.path.insert(0, str(SCRIPTS))

import rubric_authoring  # noqa: E402
from rubric_authoring import (  # noqa: E402
    AUTHORING_SCHEMA,
    AuthoringRefusal,
    assert_safe_output_target,
    build_weave_outputs,
    normalize_source,
)
from rubric_package_lib import validate_package_path  # noqa: E402


def diagnostic_codes(exc: AuthoringRefusal) -> set[str]:
    return {item["code"] for item in exc.diagnostics}


def test_explicit_markdown_normalizes_to_authoring_schema() -> None:
    contract = normalize_source(FIXTURES / "three_level_explicit.md")
    assert contract["schema"] == AUTHORING_SCHEMA
    assert contract["source"]["adapter"] == "markdown_table"
    assert [level["name"] for level in contract["rubrics"][0]["levels"]] == [
        "Excellent",
        "Capable",
        "Beginning",
    ]
    assert [level["multiplier"] for level in contract["rubrics"][0]["levels"]] == [
        1.0,
        0.7,
        0.0,
    ]
    assert {
        level["score_source"] for level in contract["rubrics"][0]["levels"]
    } == {"numeric_level_header"}
    schema = json.loads(
        (
            REPO_ROOT
            / "workspace/reference/schemas/rubrics/rubric_authoring_schema.json"
        ).read_text(encoding="utf-8")
    )
    assert not list(Draft7Validator(schema).iter_errors(contract))


def test_missing_scoring_and_weights_refuses_without_approvals() -> None:
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(FIXTURES / "missing_scoring_and_weights.md")
    assert {
        "SCORING_METADATA_REQUIRED",
        "CRITERION_WEIGHT_REQUIRED",
    }.issubset(diagnostic_codes(caught.value))


def test_explicit_fallback_approvals_are_recorded() -> None:
    contract = normalize_source(
        FIXTURES / "missing_scoring_and_weights.md",
        allow_even_spacing=True,
        allow_equal_weights=True,
    )
    assert contract["approvals"] == {
        "even_spacing": True,
        "equal_weights": True,
    }
    assert {"EVEN_SPACING_APPROVED", "EQUAL_WEIGHTS_APPROVED"} <= {
        item["code"] for item in contract["diagnostics"]
    }
    assert {
        criterion["weight_source"]
        for criterion in contract["rubrics"][0]["criteria"]
    } == {"approved_equal_weights"}


def test_authoring_json_reingest_preserves_declared_approvals(tmp_path: Path) -> None:
    first = normalize_source(
        FIXTURES / "missing_scoring_and_weights.md",
        allow_even_spacing=True,
        allow_equal_weights=True,
    )
    source = tmp_path / "normalized.json"
    source.write_text(json.dumps(first), encoding="utf-8")
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(source)
    assert {
        "SCORING_METADATA_REQUIRED",
        "CRITERION_WEIGHT_REQUIRED",
        "PRIOR_APPROVAL_OBSERVED",
    } <= diagnostic_codes(caught.value)

    second = normalize_source(
        source,
        allow_even_spacing=True,
        allow_equal_weights=True,
    )
    assert second["approvals"] == first["approvals"]
    assert second["rubrics"] == first["rubrics"]
    assert second["extensions"]["prior_approvals"] == first["approvals"]
    assert {"EVEN_SPACING_APPROVED", "EQUAL_WEIGHTS_APPROVED"} <= {
        item["code"] for item in second["diagnostics"]
    }


def test_build_is_deterministic_minimal_and_receipted(tmp_path: Path) -> None:
    first = build_weave_outputs(
        FIXTURES / "three_level_explicit.md",
        tmp_path / "first",
    )
    second = build_weave_outputs(
        FIXTURES / "three_level_explicit.md",
        tmp_path / "second",
    )
    for filename in (
        "rubric_package.zip",
        "rubrics_d2l.xml",
        "normalized_rubric_authoring.json",
        "rubric_mapping.md",
        "diagnostics.json",
        "run_receipt.json",
    ):
        assert (tmp_path / "first" / filename).read_bytes() == (
            tmp_path / "second" / filename
        ).read_bytes()
    with ZipFile(first["zip_path"]) as archive:
        assert archive.namelist() == [
            "imsmanifest.xml",
            "orgunitconfig/orgunitconfig.xml",
            "rubrics_d2l.xml",
        ]
        for info in archive.infolist():
            assert info.date_time == (1980, 1, 1, 0, 0, 0)
            assert (info.external_attr >> 16) & 0o777 == 0o644
    assert first["rubrics_xml_path"].read_bytes() == (
        first["package_dir"] / "rubrics_d2l.xml"
    ).read_bytes()
    receipt = json.loads(first["receipt_path"].read_text(encoding="utf-8"))
    assert receipt["schema"] == "coursecraft.run/1"
    assert receipt["status"] == "ok"
    assert receipt["producer"]["identity_state"] == "git"
    assert len(receipt["producer"]["commit"]) == 40
    assert receipt["producer"]["repository"] == "coursecraft-workbench"
    assert not Path(receipt["producer"]["repository"]).is_absolute()
    assert receipt["extensions"]["approvals"] == {
        "even_spacing": False,
        "equal_weights": False,
    }
    assert all(not path.endswith(("dropbox_d2l.xml", "discussion_d2l.xml")) for path in [
        artifact["path"] for artifact in receipt["emitted_files"]
    ])


def test_variable_levels_and_descriptions_survive_semantic_roundtrip(tmp_path: Path) -> None:
    result = build_weave_outputs(
        FIXTURES / "three_level_explicit.md",
        tmp_path / "out",
    )
    root = ET.parse(result["rubrics_xml_path"]).getroot()
    levels = root.findall("./rubric/criteria_groups/criteria_group/level_set/levels/level")
    assert [level.attrib["name"] for level in levels] == [
        "Excellent",
        "Capable",
        "Beginning",
    ]
    cells = root.findall(
        "./rubric/criteria_groups/criteria_group/criteria/criterion[1]/cells/cell"
    )
    assert [float(cell.attrib["cell_value"]) for cell in cells] == [60.0, 42.0, 0.0]
    assert cells[0].findtext("./description/text") == "<p>Makes a precise claim.</p>"


def test_extraction_adapter_accepts_only_consistent_numeric_grid() -> None:
    contract = normalize_source(FIXTURES / "eligible_extraction.json")
    rubric = contract["rubrics"][0]
    assert [level["multiplier"] for level in rubric["levels"]] == [1.0, 0.6, 0.0]
    assert [criterion["weight"] for criterion in rubric["criteria"]] == [60.0, 40.0]
    assert {
        criterion["weight_source"] for criterion in rubric["criteria"]
    } == {"extracted_cell_points"}

    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(FIXTURES / "inconsistent_extraction.json")
    assert "RUBRICS_ADAPTER_UNTRANSLATABLE" in diagnostic_codes(caught.value)


def test_duplicate_names_and_identifier_collisions_refuse(tmp_path: Path) -> None:
    data = {
        "rubrics": [
            {
                "name": "Same",
                "source_id": "7",
                "resource_code": "same-code",
                "levels": [
                    {"name": "High", "multiplier": 1},
                    {"name": "Low", "multiplier": 0},
                ],
                "criteria": [
                    {
                        "name": "Repeated",
                        "weight": 50,
                        "levels": {"High": "H", "Low": "L"},
                    },
                    {
                        "name": "Repeated",
                        "weight": 50,
                        "levels": {"High": "H2", "Low": "L2"},
                    },
                ],
            },
            {
                "name": "same",
                "source_id": "7",
                "resource_code": "same-code",
                "levels": [
                    {"name": "High", "multiplier": 1},
                    {"name": "Low", "multiplier": 0},
                ],
                "criteria": [
                    {
                        "name": "Only",
                        "weight": 100,
                        "levels": {"High": "H", "Low": "L"},
                    }
                ],
            },
        ]
    }
    source = tmp_path / "duplicates.json"
    source.write_text(json.dumps(data), encoding="utf-8")
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(source)
    assert {
        "DUPLICATE_RUBRIC_NAME",
        "DUPLICATE_CRITERION_NAME",
        "IDENTIFIER_COLLISION",
    } <= diagnostic_codes(caught.value)


@pytest.mark.parametrize("member", ["../outside.txt", "/absolute.txt", "extra.txt"])
def test_archive_validator_rejects_unsafe_or_unexpected_members(
    tmp_path: Path,
    member: str,
) -> None:
    archive_path = tmp_path / "unsafe.zip"
    with ZipFile(archive_path, "w", ZIP_DEFLATED) as archive:
        for required in (
            "imsmanifest.xml",
            "orgunitconfig/orgunitconfig.xml",
            "rubrics_d2l.xml",
        ):
            archive.writestr(required, b"x")
        archive.writestr(member, b"x")
    errors, _, _ = validate_package_path(archive_path)
    assert errors


def test_archive_validator_rejects_duplicate_and_symlink_members(tmp_path: Path) -> None:
    archive_path = tmp_path / "unsafe.zip"
    with ZipFile(archive_path, "w", ZIP_DEFLATED) as archive:
        archive.writestr("imsmanifest.xml", b"x")
        with pytest.warns(UserWarning, match="Duplicate name"):
            archive.writestr("imsmanifest.xml", b"y")
        archive.writestr("orgunitconfig/orgunitconfig.xml", b"x")
        info = ZipInfo("rubrics_d2l.xml")
        info.create_system = 3
        info.external_attr = 0o120777 << 16
        archive.writestr(info, b"target")
    errors, _, _ = validate_package_path(archive_path)
    assert any("duplicate" in error.lower() for error in errors)
    assert any("symlink" in error.lower() for error in errors)


def _write_straight_docx(path: Path) -> None:
    document = Document()
    document.add_heading("DOCX Synthetic Rubric", level=2)
    document.add_paragraph("Level scores: Excellent=100; Beginning=0")
    table = document.add_table(rows=2, cols=4)
    for cell, value in zip(
        table.rows[0].cells,
        ["Criterion", "Weight", "Excellent", "Beginning"],
    ):
        cell.text = value
    for cell, value in zip(
        table.rows[1].cells,
        ["Evidence", "100", "Uses evidence.", "Does not use evidence."],
    ):
        cell.text = value
    document.save(path)


def test_docx_straight_table_builds_conversion_review(tmp_path: Path) -> None:
    source = tmp_path / "rubric.docx"
    _write_straight_docx(source)
    result = build_weave_outputs(source, tmp_path / "out")
    assert result["conversion_review_path"].exists()
    contract = json.loads(result["normalized_json_path"].read_text(encoding="utf-8"))
    assert contract["source"]["adapter"] == "docx_table"


def test_docx_merged_and_nested_tables_refuse(tmp_path: Path) -> None:
    merged = tmp_path / "merged.docx"
    document = Document()
    document.add_heading("Merged", level=2)
    table = document.add_table(rows=2, cols=4)
    table.cell(0, 2).merge(table.cell(0, 3))
    document.save(merged)
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(merged, allow_even_spacing=True, allow_equal_weights=True)
    assert "DOCX_MERGED_CELL_UNSUPPORTED" in diagnostic_codes(caught.value)

    nested = tmp_path / "nested.docx"
    document = Document()
    document.add_heading("Nested", level=2)
    table = document.add_table(rows=2, cols=4)
    table.cell(1, 1).add_table(rows=1, cols=1)
    document.save(nested)
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(nested, allow_even_spacing=True, allow_equal_weights=True)
    assert "DOCX_NESTED_TABLE_UNSUPPORTED" in diagnostic_codes(caught.value)


def test_preflight_cli_is_machine_readable_and_refuses_headless_guessing() -> None:
    result = subprocess.run(
        [
            sys.executable,
            "scripts/make_rubric_package.py",
            "--input",
            str(FIXTURES / "missing_scoring_and_weights.md"),
            "--preflight",
        ],
        cwd=REPO_ROOT,
        text=True,
        capture_output=True,
        check=False,
    )
    assert result.returncode == 2
    payload = json.loads(result.stderr)
    assert payload["status"] == "error"
    assert {
        "SCORING_METADATA_REQUIRED",
        "CRITERION_WEIGHT_REQUIRED",
    } <= {item["code"] for item in payload["diagnostics"]}


def _strict_builder_contract(level_count: int) -> dict[str, object]:
    multipliers = [
        round(0.8 * (level_count - index - 1) / (level_count - 1), 6)
        for index in range(level_count)
    ]
    names = [f"Band {index + 1}" for index in range(level_count)]
    return {
        "rubrics": [
            {
                "name": f"Synthetic {level_count}-level rubric",
                "levels": [
                    {"name": name, "multiplier": multiplier}
                    for name, multiplier in zip(names, multipliers)
                ],
                "criteria": [
                    {
                        "name": "Evidence",
                        "weight": 100,
                        "levels": {
                            name: f"Authored description {index + 1}."
                            for index, name in enumerate(names)
                        },
                    }
                ],
            }
        ]
    }


@pytest.mark.parametrize("level_count", [2, 4, 5])
def test_strict_two_four_and_five_level_families(
    tmp_path: Path,
    level_count: int,
) -> None:
    source = tmp_path / f"family-{level_count}.json"
    source.write_text(json.dumps(_strict_builder_contract(level_count)), encoding="utf-8")
    contract = normalize_source(source)
    assert len(contract["rubrics"][0]["levels"]) == level_count
    assert max(level["multiplier"] for level in contract["rubrics"][0]["levels"]) == 0.8
    result = build_weave_outputs(source, tmp_path / f"out-{level_count}")
    records = rubric_authoring.rubrics_to_records(result["rubrics_xml_path"])
    assert len(records["rubrics"][0]["levels"]) == level_count


def test_overall_thresholds_are_exact_percent_multipliers(tmp_path: Path) -> None:
    source = tmp_path / "thresholds.json"
    source.write_text(
        json.dumps(
            {
                "rubrics": [
                    {
                        "name": "Threshold Rubric",
                        "levels": [{"name": "Strong"}, {"name": "Emerging"}],
                        "overall_thresholds": {"Strong": 85, "Emerging": 0},
                        "criteria": [
                            {
                                "name": "Reasoning",
                                "weight": 100,
                                "levels": {
                                    "Strong": "Strong reasoning.",
                                    "Emerging": "Emerging reasoning.",
                                },
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    contract = normalize_source(source)
    levels = contract["rubrics"][0]["levels"]
    assert [(level["multiplier"], level["score_source"]) for level in levels] == [
        (0.85, "overall_threshold_metadata"),
        (0.0, "overall_threshold_metadata"),
    ]
    assert [
        (item["name"], item["range_start_value"], item["source"])
        for item in contract["rubrics"][0]["overall_levels"]
    ] == [
        ("Emerging", 0.0, "explicit_overall_threshold"),
        ("Strong", 85.0, "explicit_overall_threshold"),
    ]
    build_weave_outputs(source, tmp_path / "out")


def test_preflight_reports_producer_owned_semantic_facts() -> None:
    contract = normalize_source(FIXTURES / "three_level_explicit.md")
    summary = rubric_authoring.preflight_summary(contract)
    rubric = summary["rubrics"][0]
    assert rubric["levels"] == [
        {
            "name": "Excellent",
            "multiplier": 1.0,
            "score_source": "numeric_level_header",
            "overall_threshold": 100.0,
            "overall_threshold_source": "derived_from_level_multiplier",
        },
        {
            "name": "Capable",
            "multiplier": 0.7,
            "score_source": "numeric_level_header",
            "overall_threshold": 70.0,
            "overall_threshold_source": "derived_from_level_multiplier",
        },
        {
            "name": "Beginning",
            "multiplier": 0.0,
            "score_source": "numeric_level_header",
            "overall_threshold": 0.0,
            "overall_threshold_source": "derived_from_level_multiplier",
        },
    ]
    assert rubric["criteria"][0] == {
        "name": "Claim",
        "weight": 60.0,
        "weight_source": "explicit_weight",
    }


@pytest.mark.parametrize(
    ("content", "code"),
    [
        ("{not json", "SOURCE_SCHEMA_UNSUPPORTED"),
        ('{"schema":"unknown/9","rubrics":[]}', "SOURCE_SCHEMA_UNSUPPORTED"),
    ],
)
def test_malformed_and_unsupported_json_refuse(
    tmp_path: Path,
    content: str,
    code: str,
) -> None:
    source = tmp_path / "bad.json"
    source.write_text(content, encoding="utf-8")
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(source)
    assert code in diagnostic_codes(caught.value)


def test_xml_escaping_survives_canonical_extractor(tmp_path: Path) -> None:
    data = _strict_builder_contract(2)
    data["rubrics"][0]["criteria"][0]["levels"]["Band 1"] = "A & B < C > D"
    source = tmp_path / "escaping.json"
    source.write_text(json.dumps(data), encoding="utf-8")
    result = build_weave_outputs(source, tmp_path / "out")
    records = rubric_authoring.rubrics_to_records(result["rubrics_xml_path"])
    assert records["rubrics"][0]["criteria"][0]["cells"][0]["description"] == "A & B < C > D"


@pytest.mark.parametrize(
    "table",
    [
        (
            "## Bad separator\n\n"
            "| Criterion | Weight | High (100) | Low (0) |\n"
            "| Criterion | Weight | High | Low |\n"
            "| One | 100 | H | L |\n"
        ),
        (
            "## Auxiliary\n\n"
            "| # | Criterion | Weight | High (100) | Low (0) |\n"
            "| --- | --- | --- | --- | --- |\n"
            "| 1 | One | 100 | H | L |\n"
        ),
    ],
)
def test_markdown_separator_and_auxiliary_columns_refuse(
    tmp_path: Path,
    table: str,
) -> None:
    source = tmp_path / "bad.md"
    source.write_text(table, encoding="utf-8")
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(source)
    assert "SOURCE_SCHEMA_UNSUPPORTED" in diagnostic_codes(caught.value)
    assert all(not code.startswith("DOCX_") for code in diagnostic_codes(caught.value))


def test_partial_weights_refuse_even_with_equal_weight_approval(tmp_path: Path) -> None:
    source = tmp_path / "partial.md"
    source.write_text(
        "## Partial\n\n"
        "| Criterion | Weight | High (100) | Low (0) |\n"
        "| --- | --- | --- | --- |\n"
        "| One | 60 | H | L |\n"
        "| Two | | H2 | L2 |\n",
        encoding="utf-8",
    )
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(source, allow_equal_weights=True)
    assert "CRITERION_WEIGHT_REQUIRED" in diagnostic_codes(caught.value)


def test_docx_ambiguous_title_nonrectangular_and_auxiliary_refuse(
    tmp_path: Path,
) -> None:
    ambiguous = tmp_path / "ambiguous.docx"
    document = Document()
    document.add_paragraph("First possible title")
    document.add_paragraph("Second possible title")
    document.add_table(rows=2, cols=4)
    document.save(ambiguous)
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(ambiguous, allow_even_spacing=True, allow_equal_weights=True)
    assert "DOCX_TABLE_AMBIGUOUS" in diagnostic_codes(caught.value)

    nonrectangular = tmp_path / "nonrectangular.docx"
    document = Document()
    document.add_paragraph("Nonrectangular")
    table = document.add_table(rows=2, cols=4)
    table.rows[1]._tr.remove(table.rows[1].cells[-1]._tc)
    document.save(nonrectangular)
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(nonrectangular, allow_even_spacing=True, allow_equal_weights=True)
    assert "DOCX_TABLE_AMBIGUOUS" in diagnostic_codes(caught.value)

    auxiliary = tmp_path / "auxiliary.docx"
    document = Document()
    document.add_paragraph("Auxiliary")
    table = document.add_table(rows=2, cols=5)
    for cell, value in zip(
        table.rows[0].cells,
        ["#", "Criterion", "Weight", "High (100)", "Low (0)"],
    ):
        cell.text = value
    document.save(auxiliary)
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(auxiliary, allow_even_spacing=True, allow_equal_weights=True)
    assert "DOCX_TABLE_AMBIGUOUS" in diagnostic_codes(caught.value)


@pytest.mark.parametrize("missing", ["rubric", "criterion", "level"])
def test_missing_structural_names_refuse(tmp_path: Path, missing: str) -> None:
    data = _strict_builder_contract(2)
    if missing == "rubric":
        data["rubrics"][0]["name"] = None
    elif missing == "criterion":
        data["rubrics"][0]["criteria"][0]["name"] = None
    else:
        data["rubrics"][0]["levels"][0]["name"] = None
    source = tmp_path / f"missing-{missing}.json"
    source.write_text(json.dumps(data), encoding="utf-8")
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(source)
    assert "SOURCE_SCHEMA_UNSUPPORTED" in diagnostic_codes(caught.value)


def test_unicode_labels_preserve_distinct_semantic_identity(tmp_path: Path) -> None:
    data = _strict_builder_contract(2)
    rubric = data["rubrics"][0]
    rubric["name"] = "分析ルーブリック"
    rubric["levels"][0]["name"] = "優秀"
    rubric["levels"][1]["name"] = "発展中"
    rubric["criteria"][0]["name"] = "証拠"
    rubric["criteria"][0]["levels"] = {
        "優秀": "十分な証拠。",
        "発展中": "証拠が必要。",
    }
    source = tmp_path / "unicode.json"
    source.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
    contract = normalize_source(source)
    assert contract["rubrics"][0]["name"] == "分析ルーブリック"
    assert [level["name"] for level in contract["rubrics"][0]["levels"]] == [
        "優秀",
        "発展中",
    ]


def test_force_guard_refuses_protected_containing_and_symlink_targets(
    tmp_path: Path,
) -> None:
    source = FIXTURES / "three_level_explicit.md"
    for protected in (
        Path("/"),
        Path("/tmp"),
        Path("/private/tmp"),
        Path(tempfile.gettempdir()),
        REPO_ROOT,
        REPO_ROOT.parent,
        Path.cwd(),
        Path.home(),
    ):
        with pytest.raises(ValueError):
            assert_safe_output_target(protected, input_path=source)

    containing = tmp_path / "containing"
    containing.mkdir()
    contained_source = containing / "source.md"
    contained_source.write_text(
        (FIXTURES / "three_level_explicit.md").read_text(encoding="utf-8"),
        encoding="utf-8",
    )
    sentinel = containing / "keep.txt"
    sentinel.write_text("keep", encoding="utf-8")
    with pytest.raises(ValueError):
        build_weave_outputs(
            contained_source,
            containing,
            force=True,
        )
    assert contained_source.exists() and sentinel.read_text(encoding="utf-8") == "keep"

    context = tmp_path / "context-target"
    context.mkdir()
    with pytest.raises(ValueError):
        assert_safe_output_target(context, input_path=source, context_dir=context)

    real = tmp_path / "real"
    real.mkdir()
    link = tmp_path / "linked-output"
    link.symlink_to(real, target_is_directory=True)
    with pytest.raises(ValueError):
        assert_safe_output_target(link, input_path=source)


def test_folder_and_zip_are_both_validated(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    calls: list[Path] = []
    original = rubric_authoring.validate_package_path

    def spy(path: Path):
        calls.append(path)
        return original(path)

    monkeypatch.setattr(rubric_authoring, "validate_package_path", spy)
    result = build_weave_outputs(
        FIXTURES / "three_level_explicit.md",
        tmp_path / "out",
    )
    assert any(path.name == "package" for path in calls)
    assert any(path.suffix == ".zip" for path in calls)


def test_malformed_folder_zip_symlink_and_oversize_fail_closed(tmp_path: Path) -> None:
    folder = tmp_path / "package"
    (folder / "orgunitconfig").mkdir(parents=True)
    (folder / "imsmanifest.xml").write_text("<manifest>", encoding="utf-8")
    (folder / "orgunitconfig" / "orgunitconfig.xml").write_text("<orgunit/>", encoding="utf-8")
    (folder / "rubrics_d2l.xml").write_text("<rubrics/>", encoding="utf-8")
    errors, _, _ = validate_package_path(folder)
    assert any("malformed XML" in error for error in errors)

    archive = tmp_path / "malformed.zip"
    with ZipFile(archive, "w", ZIP_DEFLATED) as zipped:
        zipped.writestr("imsmanifest.xml", "<manifest>")
        zipped.writestr("orgunitconfig/orgunitconfig.xml", "<orgunit/>")
        zipped.writestr("rubrics_d2l.xml", "<rubrics/>")
    errors, _, _ = validate_package_path(archive)
    assert any("malformed XML" in error for error in errors)

    good_target = tmp_path / "real-rubrics.xml"
    good_target.write_text("<rubrics schemaversion='v2011'/>", encoding="utf-8")
    (folder / "imsmanifest.xml").write_text("<manifest/>", encoding="utf-8")
    (folder / "rubrics_d2l.xml").unlink()
    (folder / "rubrics_d2l.xml").symlink_to(good_target)
    errors, _, _ = validate_package_path(folder)
    assert any("symlink" in error for error in errors)

    (folder / "rubrics_d2l.xml").unlink()
    (folder / "rubrics_d2l.xml").write_bytes(b"x" * (10 * 1024 * 1024 + 1))
    errors, _, _ = validate_package_path(folder)
    assert any("10 MiB" in error for error in errors)


def test_receipt_artifacts_hashes_generic_identity_and_no_path_leakage(
    tmp_path: Path,
) -> None:
    result = build_weave_outputs(
        FIXTURES / "three_level_explicit.md",
        tmp_path / "out",
    )
    receipt = json.loads(result["receipt_path"].read_text(encoding="utf-8"))
    run_schema = json.loads(
        (
            REPO_ROOT
            / "workspace/reference/schemas/course/run_identity_schema.json"
        ).read_text(encoding="utf-8")
    )
    assert not list(Draft7Validator(run_schema).iter_errors(receipt))
    for artifact in receipt["emitted_files"]:
        path = result["output_dir"] / artifact["path"]
        assert artifact["bytes"] == path.stat().st_size
        assert artifact["sha256"] == hashlib.sha256(path.read_bytes()).hexdigest()
    text_outputs = [
        path
        for path in result["output_dir"].rglob("*")
        if path.is_file() and path.suffix in {".json", ".md", ".xml"}
    ]
    combined = "\n".join(path.read_text(encoding="utf-8") for path in text_outputs)
    assert str(REPO_ROOT) not in combined
    assert str(Path.home()) not in combined
    assert "three_level_explicit" not in combined
    contract = json.loads(result["normalized_json_path"].read_text(encoding="utf-8"))
    assert contract["source"]["label"] == "rubric-source"
    manifest = ET.parse(result["package_dir"] / "imsmanifest.xml")
    namespace = {"imsmd": "http://www.imsglobal.org/xsd/imsmd_rootv1p2p1"}
    assert manifest.findtext(".//imsmd:title/imsmd:langstring", namespaces=namespace) == "Rubric Package"


def test_deterministic_identifiers_no_activity_resources_and_canonical_roundtrip(
    tmp_path: Path,
) -> None:
    first = build_weave_outputs(
        FIXTURES / "three_level_explicit.md",
        tmp_path / "first",
    )
    second = build_weave_outputs(
        FIXTURES / "three_level_explicit.md",
        tmp_path / "second",
    )
    first_records = rubric_authoring.rubrics_to_records(first["rubrics_xml_path"])
    second_records = rubric_authoring.rubrics_to_records(second["rubrics_xml_path"])
    assert first_records["rubrics"][0]["id"] == second_records["rubrics"][0]["id"] == "1"
    assert first_records["rubrics"][0]["resource_code"] == second_records["rubrics"][0]["resource_code"]
    assert [level["level_id"] for level in first_records["rubrics"][0]["levels"]] == [
        level["level_id"] for level in second_records["rubrics"][0]["levels"]
    ]
    assert first_records["rubrics"][0]["criteria"][0]["cells"][0]["description"] == (
        "Makes a precise claim."
    )
    manifest = ET.parse(first["package_dir"] / "imsmanifest.xml").getroot()
    material_types = {
        value
        for resource in manifest.iter()
        for key, value in resource.attrib.items()
        if key.endswith("material_type")
    }
    assert material_types == {"orgunitconfig", "d2lrubrics"}
    all_names = [path.name for path in first["package_dir"].rglob("*") if path.is_file()]
    assert not any(name.startswith(("dropbox_d2l", "discussion_d2l", "quiz_d2l")) for name in all_names)


@pytest.mark.parametrize(
    ("mutation", "expected_code"),
    [
        ("invalid_level", "SCORING_METADATA_REQUIRED"),
        ("invalid_weight", "CRITERION_WEIGHT_REQUIRED"),
        ("partial_threshold", "SCORING_METADATA_REQUIRED"),
    ],
)
def test_invalid_explicit_data_never_falls_through(
    tmp_path: Path,
    mutation: str,
    expected_code: str,
) -> None:
    data = _strict_builder_contract(2)
    rubric = data["rubrics"][0]
    if mutation == "invalid_level":
        rubric["levels"][0]["multiplier"] = "not-a-score"
        rubric["levels"][1].pop("multiplier")
    elif mutation == "invalid_weight":
        rubric["criteria"][0]["weight"] = "not-a-weight"
    else:
        for level in rubric["levels"]:
            level.pop("multiplier")
        rubric["overall_thresholds"] = {"Band 1": 80}
    source = tmp_path / f"{mutation}.json"
    source.write_text(json.dumps(data), encoding="utf-8")
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(
            source,
            allow_even_spacing=True,
            allow_equal_weights=True,
        )
    assert expected_code in diagnostic_codes(caught.value)


def test_input_path_fail_closed_and_content_caps(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    missing = tmp_path / "missing.md"
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(missing)
    assert "SOURCE_PATH_MISSING" in diagnostic_codes(caught.value)

    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(tmp_path)
    assert "SOURCE_PATH_UNSAFE" in diagnostic_codes(caught.value)

    real = tmp_path / "real.md"
    real.write_text("## X", encoding="utf-8")
    linked = tmp_path / "linked.md"
    linked.symlink_to(real)
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(linked)
    assert "SOURCE_PATH_UNSAFE" in diagnostic_codes(caught.value)

    original_open = Path.open

    def denied(path: Path, *args, **kwargs):
        if path == real:
            raise PermissionError("private-path-must-not-leak")
        return original_open(path, *args, **kwargs)

    monkeypatch.setattr(Path, "open", denied)
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(real)
    assert "SOURCE_UNREADABLE" in diagnostic_codes(caught.value)
    monkeypatch.setattr(Path, "open", original_open)

    oversized = tmp_path / "oversized.md"
    with oversized.open("wb") as handle:
        handle.truncate(rubric_authoring.MAX_TEXT_SOURCE_BYTES + 1)
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(oversized)
    assert "SOURCE_SIZE_LIMIT" in diagnostic_codes(caught.value)


@pytest.mark.parametrize(
    "members",
    [
        {
            "word/document.xml": b"<w:document/>",
            "../escape.xml": b"x",
        },
        {
            "word/styles.xml": b"x",
        },
    ],
)
def test_docx_archive_preflight_refuses_unsafe_shapes(
    tmp_path: Path,
    members: dict[str, bytes],
) -> None:
    source = tmp_path / "unsafe.docx"
    with ZipFile(source, "w", ZIP_DEFLATED) as archive:
        for name, content in members.items():
            archive.writestr(name, content)
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(source)
    assert "DOCX_ARCHIVE_UNSAFE" in diagnostic_codes(caught.value)


def test_private_path_tokens_do_not_leak_and_authored_text_stays_out_of_receipt(
    tmp_path: Path,
) -> None:
    path_token = "PRIVATE_PATH_TOKEN_7f83"
    rubric_token = "PRIVATE_RUBRIC_TOKEN_91ac"
    private_dir = tmp_path / path_token
    private_dir.mkdir()
    source = private_dir / f"{path_token}.json"
    data = _strict_builder_contract(2)
    data["rubrics"][0]["name"] = rubric_token
    source.write_text(json.dumps(data), encoding="utf-8")
    result = build_weave_outputs(source, tmp_path / "out")
    normalized = result["normalized_json_path"].read_text(encoding="utf-8")
    receipt = result["receipt_path"].read_text(encoding="utf-8")
    mapping = result["mapping_path"].read_text(encoding="utf-8")
    diagnostics = result["diagnostics_path"].read_text(encoding="utf-8")
    assert rubric_token in normalized
    assert rubric_token.replace("_", r"\_") in mapping
    assert rubric_token not in receipt
    for emitted in (normalized, receipt, mapping, diagnostics):
        assert path_token not in emitted

    cli = subprocess.run(
        [
            sys.executable,
            "scripts/make_rubric_package.py",
            "--input",
            str(source),
            "--output-dir",
            str(Path.home()),
            "--force",
        ],
        cwd=REPO_ROOT,
        text=True,
        capture_output=True,
        check=False,
    )
    assert cli.returncode == 1
    assert path_token not in cli.stderr


@pytest.mark.parametrize(
    "metadata",
    [
        "- Level scores: High=100; High=50; Low=0\n",
        "- Level scores: High=100; Low=0\n- Level scores: High=90; Low=0\n",
        "- Level scores: High=100; Low=0\n- Level values: High=90; Low=0\n",
    ],
)
def test_markdown_duplicate_or_competing_score_metadata_refuses(
    tmp_path: Path,
    metadata: str,
) -> None:
    source = tmp_path / "ambiguous-scores.md"
    source.write_text(
        "## Ambiguous\n\n"
        + metadata
        + "\n| Criterion | Weight | High | Low |\n"
        "| --- | --- | --- | --- |\n"
        "| Evidence | 100 | Strong. | Weak. |\n",
        encoding="utf-8",
    )
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(source, allow_even_spacing=True)
    assert "SCORING_METADATA_REQUIRED" in diagnostic_codes(caught.value)


@pytest.mark.parametrize(
    "headers",
    [
        ["High 100/90", "Mid 50", "Low 0"],
        ["High 100", "Mid", "Low 0"],
        ["High 100 90", "Mid 50", "Low 0"],
    ],
)
def test_markdown_ambiguous_or_partial_numeric_headers_refuse(
    tmp_path: Path,
    headers: list[str],
) -> None:
    source = tmp_path / "ambiguous-headers.md"
    source.write_text(
        "## Ambiguous\n\n"
        f"| Criterion | Weight | {headers[0]} | {headers[1]} | {headers[2]} |\n"
        "| --- | --- | --- | --- | --- |\n"
        "| Evidence | 100 | Strong. | Developing. | Weak. |\n",
        encoding="utf-8",
    )
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(source, allow_even_spacing=True)
    assert "SCORING_METADATA_REQUIRED" in diagnostic_codes(caught.value)


def test_duplicate_json_keys_invalid_level_types_and_unknown_fields_refuse(
    tmp_path: Path,
) -> None:
    duplicate = tmp_path / "duplicate-key.json"
    duplicate.write_text(
        '{"rubrics":[],"rubrics":[{"name":"private","levels":[],"criteria":[]}]}',
        encoding="utf-8",
    )
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(duplicate)
    assert "SOURCE_SCHEMA_UNSUPPORTED" in diagnostic_codes(caught.value)

    invalid_levels = tmp_path / "invalid-levels.json"
    invalid_levels.write_text(
        json.dumps(
            {
                "rubrics": [
                    {
                        "name": "Invalid",
                        "levels": [42, 99],
                        "criteria": [
                            {
                                "name": "Evidence",
                                "levels": {"Level 1": "One.", "Level 2": "Two."},
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(
            invalid_levels,
            allow_even_spacing=True,
            allow_equal_weights=True,
        )
    assert "SOURCE_SCHEMA_UNSUPPORTED" in diagnostic_codes(caught.value)

    unknown = tmp_path / "unknown-fields.json"
    unknown_data = _strict_builder_contract(2)
    unknown_data["activities"] = [{"type": "assignment"}]
    unknown_data["rubrics"][0]["untranslatable"] = "must not disappear"
    unknown.write_text(json.dumps(unknown_data), encoding="utf-8")
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(unknown)
    assert "SOURCE_SCHEMA_UNSUPPORTED" in diagnostic_codes(caught.value)


def test_docx_repeated_and_duplicate_score_metadata_refuses(tmp_path: Path) -> None:
    for filename, metadata in (
        (
            "repeated.docx",
            ["Level scores: High=100; Low=0", "Level scores: High=90; Low=0"],
        ),
        (
            "duplicate-pair.docx",
            ["Level scores: High=100; High=50; Low=0"],
        ),
    ):
        source = tmp_path / filename
        document = Document()
        document.add_heading("Private title", level=2)
        for paragraph in metadata:
            document.add_paragraph(paragraph)
        table = document.add_table(rows=2, cols=4)
        for cell, value in zip(
            table.rows[0].cells,
            ["Criterion", "Weight", "High", "Low"],
        ):
            cell.text = value
        for cell, value in zip(
            table.rows[1].cells,
            ["Evidence", "100", "Strong.", "Weak."],
        ):
            cell.text = value
        document.save(source)
        with pytest.raises(AuthoringRefusal) as caught:
            normalize_source(source, allow_even_spacing=True)
        assert {
            "DOCX_TABLE_AMBIGUOUS",
            "SCORING_METADATA_REQUIRED",
        } & diagnostic_codes(caught.value)


def test_declared_contract_schema_failures_and_error_diagnostics_refuse(
    tmp_path: Path,
) -> None:
    contract = normalize_source(FIXTURES / "three_level_explicit.md")

    malformed = json.loads(json.dumps(contract))
    malformed["rubrics"] = "not-an-array"
    malformed_source = tmp_path / "schema-invalid.json"
    malformed_source.write_text(json.dumps(malformed), encoding="utf-8")
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(malformed_source)
    assert "SOURCE_SCHEMA_UNSUPPORTED" in diagnostic_codes(caught.value)

    unresolved = json.loads(json.dumps(contract))
    unresolved["diagnostics"].append(
        {
            "id": "diag-0001",
            "code": "PRIVATE_ERROR",
            "severity": "error",
            "message": "private authored diagnostic",
            "location": "source",
            "remediation": "repair",
            "extensions": {},
        }
    )
    unresolved_source = tmp_path / "unresolved.json"
    unresolved_source.write_text(json.dumps(unresolved), encoding="utf-8")
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(unresolved_source)
    assert "SOURCE_SCHEMA_UNSUPPORTED" in diagnostic_codes(caught.value)
    assert all(
        "private authored diagnostic" not in item["message"]
        for item in caught.value.diagnostics
    )

    extraction = json.loads(
        (FIXTURES / "eligible_extraction.json").read_text(encoding="utf-8")
    )
    extraction["rubrics"] = "PRIVATE_INVALID_EXTRACTION_SHAPE"
    extraction_source = tmp_path / "invalid-extraction.json"
    extraction_source.write_text(json.dumps(extraction), encoding="utf-8")
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(extraction_source)
    assert "RUBRICS_ADAPTER_UNTRANSLATABLE" in diagnostic_codes(caught.value)
    assert all(
        "PRIVATE_INVALID_EXTRACTION_SHAPE" not in item["message"]
        for item in caught.value.diagnostics
    )


@pytest.mark.parametrize("mutation", ["empty_level_id", "duplicate_cell", "wrong_level_name"])
def test_extraction_adapter_requires_exact_nonempty_cell_level_join(
    tmp_path: Path,
    mutation: str,
) -> None:
    data = json.loads(
        (FIXTURES / "eligible_extraction.json").read_text(encoding="utf-8")
    )
    rubric = data["rubrics"][0]
    if mutation == "empty_level_id":
        rubric["levels"][0]["level_id"] = ""
        for criterion in rubric["criteria"]:
            criterion["cells"][0]["level_id"] = ""
    elif mutation == "duplicate_cell":
        rubric["criteria"][0]["cells"][1]["level_id"] = rubric["levels"][0]["level_id"]
        rubric["criteria"][0]["cells"][1]["level_name"] = rubric["levels"][0]["name"]
    else:
        rubric["criteria"][0]["cells"][0]["level_name"] = "Wrong private label"
    source = tmp_path / f"{mutation}.json"
    source.write_text(json.dumps(data), encoding="utf-8")
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(source)
    assert "RUBRICS_ADAPTER_UNTRANSLATABLE" in diagnostic_codes(caught.value)


def test_escaped_markdown_pipes_roundtrip_and_mapping_escape(tmp_path: Path) -> None:
    source = tmp_path / "escaped.md"
    source.write_text(
        "## Rubric \\| Name\n\n"
        "| Criterion | Weight | High \\| Honors (100) | Low (0) |\n"
        "| --- | --- | --- | --- |\n"
        "| Evidence \\| Reasoning | 100 | Uses A \\| B. | Omits A \\| B. |\n",
        encoding="utf-8",
    )
    result = build_weave_outputs(source, tmp_path / "out")
    records = rubric_authoring.rubrics_to_records(result["rubrics_xml_path"])
    rubric = records["rubrics"][0]
    assert rubric["name"] == r"Rubric \| Name"
    assert rubric["levels"][0]["name"] == "High | Honors"
    assert rubric["criteria"][0]["name"] == "Evidence | Reasoning"
    assert rubric["criteria"][0]["cells"][0]["description"] == "Uses A | B."
    mapping = result["mapping_path"].read_text(encoding="utf-8")
    assert r"High \| Honors" in mapping
    assert r"Evidence \| Reasoning" in mapping


def test_force_replacement_is_transactional_for_context_and_roundtrip_failures(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    output = tmp_path / "out"
    build_weave_outputs(FIXTURES / "three_level_explicit.md", output)
    sentinel = output / "KEEP.md"
    sentinel.write_text("keep", encoding="utf-8")

    with pytest.raises(AuthoringRefusal):
        build_weave_outputs(
            FIXTURES / "three_level_explicit.md",
            output,
            cli_overrides={"identifier": "res_rubrics"},
            force=True,
        )
    assert sentinel.read_text(encoding="utf-8") == "keep"

    monkeypatch.setattr(
        rubric_authoring,
        "_semantic_roundtrip",
        lambda _contract, _xml: ["synthetic failure"],
    )
    with pytest.raises(AuthoringRefusal):
        build_weave_outputs(
            FIXTURES / "three_level_explicit.md",
            output,
            force=True,
        )
    assert sentinel.read_text(encoding="utf-8") == "keep"
    assert not list(tmp_path.glob(".out.candidate-*"))
    assert not (tmp_path / ".out.previous").exists()


def test_validator_requires_cell_level_bijection_for_folder_and_zip(
    tmp_path: Path,
) -> None:
    result = build_weave_outputs(
        FIXTURES / "three_level_explicit.md",
        tmp_path / "built",
    )
    xml_path = result["package_dir"] / "rubrics_d2l.xml"
    root = ET.parse(xml_path)
    cells = root.findall(
        "./rubric/criteria_groups/criteria_group/criteria/criterion[1]/cells/cell"
    )
    cells[1].set("level_id", cells[0].attrib["level_id"])
    root.write(xml_path, encoding="utf-8", xml_declaration=True)
    folder_errors, _, _ = validate_package_path(result["package_dir"])
    assert any("exactly once" in error for error in folder_errors)

    mutated_zip = tmp_path / "mutated.zip"
    rubric_authoring.zip_package(result["package_dir"], mutated_zip)
    zip_errors, _, _ = validate_package_path(mutated_zip)
    assert any("exactly once" in error for error in zip_errors)


def test_malformed_and_encrypted_zip_errors_are_content_minimized(tmp_path: Path) -> None:
    private_token = "PRIVATE_BAD_TOKEN_9922"
    malformed = tmp_path / f"{private_token}.zip"
    malformed.write_bytes(b"not a zip")
    errors, _, _ = validate_package_path(malformed)
    assert errors and private_token not in json.dumps(errors)
    cli = subprocess.run(
        [
            sys.executable,
            "scripts/validate_rubric_package.py",
            str(malformed),
        ],
        cwd=REPO_ROOT,
        text=True,
        capture_output=True,
        check=False,
    )
    assert cli.returncode == 1
    assert private_token not in cli.stdout + cli.stderr
    assert "Traceback" not in cli.stdout + cli.stderr

    encrypted = tmp_path / "encrypted.zip"
    with ZipFile(encrypted, "w", ZIP_DEFLATED) as archive:
        for name in (
            "imsmanifest.xml",
            "orgunitconfig/orgunitconfig.xml",
            "rubrics_d2l.xml",
        ):
            archive.writestr(name, b"x")
    payload = bytearray(encrypted.read_bytes())
    for signature, flag_offset in ((b"PK\x03\x04", 6), (b"PK\x01\x02", 8)):
        cursor = 0
        while True:
            cursor = payload.find(signature, cursor)
            if cursor < 0:
                break
            flags = int.from_bytes(
                payload[cursor + flag_offset : cursor + flag_offset + 2], "little"
            )
            payload[cursor + flag_offset : cursor + flag_offset + 2] = (
                flags | 1
            ).to_bytes(2, "little")
            cursor += 4
    encrypted.write_bytes(payload)
    errors, _, _ = validate_package_path(encrypted)
    assert errors == ["Archive contains encrypted members."]


def test_cli_refusal_and_success_paths_are_private_and_relative(tmp_path: Path) -> None:
    private_rubric = "PRIVATE_RUBRIC_TOKEN_4477"
    duplicate_data = _strict_builder_contract(2)
    second = json.loads(json.dumps(duplicate_data["rubrics"][0]))
    duplicate_data["rubrics"][0]["name"] = private_rubric
    second["name"] = private_rubric
    duplicate_data["rubrics"].append(second)
    duplicate_source = tmp_path / "duplicates.json"
    duplicate_source.write_text(json.dumps(duplicate_data), encoding="utf-8")
    refused = subprocess.run(
        [
            sys.executable,
            "scripts/make_rubric_package.py",
            "--input",
            str(duplicate_source),
            "--preflight",
        ],
        cwd=REPO_ROOT,
        text=True,
        capture_output=True,
        check=False,
    )
    assert refused.returncode == 2
    assert private_rubric not in refused.stderr

    private_output = "PRIVATE_OUTPUT_TOKEN_8821"
    output = tmp_path / private_output
    success = subprocess.run(
        [
            sys.executable,
            "scripts/make_rubric_package.py",
            "--input",
            str(FIXTURES / "three_level_explicit.md"),
            "--output-dir",
            str(output),
        ],
        cwd=REPO_ROOT,
        text=True,
        capture_output=True,
        check=False,
    )
    assert success.returncode == 0, success.stderr
    assert private_output not in success.stdout
    assert "output_dir=." in success.stdout
    assert "zip_path=rubric_package.zip" in success.stdout


def test_receipt_context_privacy_identity_provenance_and_lineage(tmp_path: Path) -> None:
    private_title = "PRIVATE_TITLE_TOKEN_5511"
    private_keyword = "PRIVATE_KEYWORD_TOKEN_6622"
    first = build_weave_outputs(
        FIXTURES / "three_level_explicit.md",
        tmp_path / "first",
        source_label="Label A",
        cli_overrides={"title": private_title, "keyword": private_keyword},
    )
    second = build_weave_outputs(
        FIXTURES / "three_level_explicit.md",
        tmp_path / "second",
        source_label="Label B",
        cli_overrides={"title": private_title, "keyword": private_keyword},
    )
    first_receipt_text = first["receipt_path"].read_text(encoding="utf-8")
    assert private_title not in first_receipt_text
    assert private_keyword not in first_receipt_text
    first_receipt = json.loads(first_receipt_text)
    second_receipt = json.loads(second["receipt_path"].read_text(encoding="utf-8"))
    assert first_receipt["run_id"] != second_receipt["run_id"]
    assert len(first_receipt["parameters"]["context_sha256"]) == 64
    assert {
        "rubric_authoring.py",
        "rubric_package_lib.py",
        "extract_rubrics_to_workbook.py",
    } <= set(first_receipt["producer"]["extensions"]["code_digests"])
    assert {
        "rubric_authoring_schema.json",
        "run_identity_schema.json",
        "rubrics_schema.json",
    } <= set(first_receipt["producer"]["extensions"]["schema_digests"])

    first_contract = json.loads(
        first["normalized_json_path"].read_text(encoding="utf-8")
    )
    reingest = tmp_path / "reingest.json"
    reingest.write_text(json.dumps(first_contract), encoding="utf-8")
    second_contract = normalize_source(reingest)
    assert first_contract["source"]["sha256"] in {
        item["sha256"]
        for item in second_contract["extensions"]["source_lineage"]
    }

    extraction = build_weave_outputs(
        FIXTURES / "eligible_extraction.json",
        tmp_path / "extraction",
    )
    extraction_receipt = json.loads(
        extraction["receipt_path"].read_text(encoding="utf-8")
    )
    extraction_contract = next(
        item
        for item in extraction_receipt["contracts"]
        if item["schema"] == "coursecraft.rubrics/1"
    )
    assert extraction_contract["extensions"]["role"] == "input_contract"
    assert extraction_contract["extensions"]["input_document_sha256"] == (
        hashlib.sha256(
            (FIXTURES / "eligible_extraction.json").read_bytes()
        ).hexdigest()
    )
    assert extraction_contract["sha256"] == hashlib.sha256(
        (
            REPO_ROOT
            / "workspace/reference/schemas/rubrics/rubrics_schema.json"
        ).read_bytes()
    ).hexdigest()


def test_canonical_equivalent_unicode_names_collide(tmp_path: Path) -> None:
    data = _strict_builder_contract(2)
    first = data["rubrics"][0]
    second = json.loads(json.dumps(first))
    first["name"] = "Café"
    second["name"] = "Cafe\u0301"
    data["rubrics"].append(second)
    source = tmp_path / "unicode-collision.json"
    source.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(source)
    assert "DUPLICATE_RUBRIC_NAME" in diagnostic_codes(caught.value)


def _write_metadata_docx(path: Path, score_metadata: str) -> None:
    document = Document()
    document.add_heading("Metadata Join", level=2)
    document.add_paragraph(f"Level scores: {score_metadata}")
    table = document.add_table(rows=2, cols=4)
    for cell, value in zip(
        table.rows[0].cells,
        ["Criterion", "Weight", "Café", "Low"],
    ):
        cell.text = value
    for cell, value in zip(
        table.rows[1].cells,
        ["Evidence", "100", "Strong.", "Weak."],
    ):
        cell.text = value
    document.save(path)


@pytest.mark.parametrize("adapter_kind", ["markdown", "docx"])
@pytest.mark.parametrize(
    "score_metadata",
    [
        "Excellent=100; Beginning=0",
        "CAFE\u0301=100; Beginning=0",
    ],
    ids=["no-level-joins", "partial-level-join"],
)
def test_present_nonjoining_score_metadata_blocks_fallback(
    tmp_path: Path,
    adapter_kind: str,
    score_metadata: str,
) -> None:
    source = tmp_path / f"join.{adapter_kind if adapter_kind == 'markdown' else 'docx'}"
    if adapter_kind == "markdown":
        source.write_text(
            "## Metadata Join\n\n"
            f"- Level scores: {score_metadata}\n\n"
            "| Criterion | Weight | Café | Low |\n"
            "| --- | --- | --- | --- |\n"
            "| Evidence | 100 | Strong. | Weak. |\n",
            encoding="utf-8",
        )
    else:
        _write_metadata_docx(source, score_metadata)
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(source, allow_even_spacing=True)
    assert "SCORING_METADATA_REQUIRED" in diagnostic_codes(caught.value)
    assert "EVEN_SPACING_APPROVED" not in diagnostic_codes(caught.value)


@pytest.mark.parametrize("adapter_kind", ["markdown", "docx"])
def test_score_metadata_joins_by_unicode_case_identity(
    tmp_path: Path,
    adapter_kind: str,
) -> None:
    score_metadata = "CAFE\u0301=100; LOW=0"
    source = tmp_path / f"unicode-join.{adapter_kind if adapter_kind == 'markdown' else 'docx'}"
    if adapter_kind == "markdown":
        source.write_text(
            "## Metadata Join\n\n"
            f"- Level scores: {score_metadata}\n\n"
            "| Criterion | Weight | Café | Low |\n"
            "| --- | --- | --- | --- |\n"
            "| Evidence | 100 | Strong. | Weak. |\n",
            encoding="utf-8",
        )
    else:
        _write_metadata_docx(source, score_metadata)
    contract = normalize_source(source, allow_even_spacing=True)
    assert [
        (item["multiplier"], item["score_source"])
        for item in contract["rubrics"][0]["levels"]
    ] == [
        (1.0, "explicit_level_metadata"),
        (0.0, "explicit_level_metadata"),
    ]
    assert "EVEN_SPACING_APPROVED" not in {
        item["code"] for item in contract["diagnostics"]
    }


def _eligible_extraction() -> dict:
    return json.loads(
        (FIXTURES / "eligible_extraction.json").read_text(encoding="utf-8")
    )


def _set_extracted_rubric_identity(
    rubric: dict,
    *,
    rubric_id: str,
    resource_code: str,
    name: str,
) -> None:
    rubric["id"] = rubric_id
    rubric["resource_code"] = resource_code
    rubric["name"] = name
    rubric["attributes"]["id"] = rubric_id
    rubric["attributes"]["resource_code"] = resource_code
    rubric["attributes"]["name"] = name


@pytest.mark.parametrize(
    "mutation",
    [
        "top_diagnostic",
        "group_diagnostic",
        "unsupported_scoring_method",
        "rubric_description",
        "overall_description",
        "overall_feedback",
        "unsupported_attribute",
        "attribute_state",
        "attribute_identity_mismatch",
        "empty_rubric_id",
        "empty_resource_code",
        "empty_rubric_name",
        "duplicate_rubric_id",
        "duplicate_resource_code",
        "duplicate_rubric_name",
        "empty_level_id",
        "duplicate_level_id",
        "empty_level_name",
        "duplicate_level_name",
        "missing_level_sort",
        "duplicate_level_sort",
        "unordered_level_sort",
        "empty_criterion_name",
        "duplicate_criterion_name",
        "missing_criterion_sort",
        "duplicate_criterion_sort",
        "unordered_criterion_sort",
        "cell_id_mismatch",
        "cell_name_mismatch",
        "points_raw_disagreement",
        "nonfinite_points",
        "maxima_not_100",
        "inconsistent_multipliers",
        "nonunique_multipliers",
        "out_of_range_multipliers",
        "overall_cardinality",
        "overall_identity",
        "overall_order",
        "duplicate_overall_name",
        "duplicate_overall_sort",
        "overall_sort_mismatch",
        "duplicate_threshold",
        "out_of_range_threshold",
        "threshold_multiplier_mismatch",
        "score_band_mismatch",
    ],
)
def test_extraction_adapter_refuses_every_unpreserved_semantic_variant(
    tmp_path: Path,
    mutation: str,
) -> None:
    data = _eligible_extraction()
    rubric = data["rubrics"][0]
    levels = rubric["levels"]
    criteria = rubric["criteria"]
    overall = rubric["overall_levels"]

    if mutation in {"top_diagnostic", "group_diagnostic"}:
        data["diagnostics"] = ["private diagnostic"]
    elif mutation == "unsupported_scoring_method":
        rubric["scoring_method"] = "999"
    elif mutation == "rubric_description":
        rubric["description"] = "private rubric prose"
    elif mutation == "overall_description":
        overall[0]["description"] = "private overall prose"
    elif mutation == "overall_feedback":
        overall[0]["feedback"] = "private feedback"
    elif mutation == "unsupported_attribute":
        rubric["attributes"]["private_state"] = "1"
    elif mutation == "attribute_state":
        rubric["attributes"]["state"] = "1"
    elif mutation == "attribute_identity_mismatch":
        rubric["attributes"]["name"] = "Different"
    elif mutation == "empty_rubric_id":
        _set_extracted_rubric_identity(
            rubric,
            rubric_id="",
            resource_code=rubric["resource_code"],
            name=rubric["name"],
        )
    elif mutation == "empty_resource_code":
        _set_extracted_rubric_identity(
            rubric,
            rubric_id=rubric["id"],
            resource_code="",
            name=rubric["name"],
        )
    elif mutation == "empty_rubric_name":
        _set_extracted_rubric_identity(
            rubric,
            rubric_id=rubric["id"],
            resource_code=rubric["resource_code"],
            name="",
        )
    elif mutation.startswith("duplicate_rubric") or mutation == "duplicate_resource_code":
        second = json.loads(json.dumps(rubric))
        _set_extracted_rubric_identity(
            second,
            rubric_id="8",
            resource_code="SYNTHETIC-RUBRIC-8",
            name="Second Synthetic Rubric",
        )
        if mutation == "duplicate_rubric_id":
            _set_extracted_rubric_identity(
                second,
                rubric_id=rubric["id"],
                resource_code=second["resource_code"],
                name=second["name"],
            )
        elif mutation == "duplicate_resource_code":
            _set_extracted_rubric_identity(
                second,
                rubric_id=second["id"],
                resource_code=rubric["resource_code"],
                name=second["name"],
            )
        else:
            _set_extracted_rubric_identity(
                second,
                rubric_id=second["id"],
                resource_code=second["resource_code"],
                name="synthetic extracted rubric",
            )
        data["rubrics"].append(second)
    elif mutation == "empty_level_id":
        levels[0]["level_id"] = ""
    elif mutation == "duplicate_level_id":
        levels[1]["level_id"] = levels[0]["level_id"]
    elif mutation == "empty_level_name":
        levels[0]["name"] = ""
    elif mutation == "duplicate_level_name":
        levels[1]["name"] = levels[0]["name"].casefold()
    elif mutation == "missing_level_sort":
        levels[0].pop("sort_order")
    elif mutation == "duplicate_level_sort":
        levels[1]["sort_order"] = levels[0]["sort_order"]
    elif mutation == "unordered_level_sort":
        levels[0]["sort_order"], levels[1]["sort_order"] = (
            levels[1]["sort_order"],
            levels[0]["sort_order"],
        )
    elif mutation == "empty_criterion_name":
        criteria[0]["name"] = ""
    elif mutation == "duplicate_criterion_name":
        criteria[1]["name"] = criteria[0]["name"].casefold()
    elif mutation == "missing_criterion_sort":
        criteria[0].pop("sort_order")
    elif mutation == "duplicate_criterion_sort":
        criteria[1]["sort_order"] = criteria[0]["sort_order"]
    elif mutation == "unordered_criterion_sort":
        criteria[0]["sort_order"], criteria[1]["sort_order"] = (
            criteria[1]["sort_order"],
            criteria[0]["sort_order"],
        )
    elif mutation == "cell_id_mismatch":
        criteria[0]["cells"][1]["level_id"] = levels[0]["level_id"]
    elif mutation == "cell_name_mismatch":
        criteria[0]["cells"][0]["level_name"] = "Wrong"
    elif mutation == "points_raw_disagreement":
        criteria[0]["cells"][0]["points_raw"] = "59"
    elif mutation == "nonfinite_points":
        criteria[0]["cells"][0]["points"] = float("inf")
    elif mutation == "maxima_not_100":
        for cell, value in zip(criteria[1]["cells"], (30, 18, 0)):
            cell["points"] = value
            cell["points_raw"] = str(value)
    elif mutation == "inconsistent_multipliers":
        criteria[1]["cells"][1]["points"] = 20
        criteria[1]["cells"][1]["points_raw"] = "20"
    elif mutation == "nonunique_multipliers":
        criteria[1]["cells"][1]["points"] = 40
        criteria[1]["cells"][1]["points_raw"] = "40"
    elif mutation == "out_of_range_multipliers":
        criteria[1]["cells"][2]["points"] = -1
        criteria[1]["cells"][2]["points_raw"] = "-1"
    elif mutation == "overall_cardinality":
        overall.pop()
    elif mutation == "overall_identity":
        overall[1]["name"] = "Different"
    elif mutation == "overall_order":
        overall[0], overall[1] = overall[1], overall[0]
    elif mutation == "duplicate_overall_name":
        overall[1]["name"] = overall[0]["name"].casefold()
    elif mutation == "duplicate_overall_sort":
        overall[1]["sort_order"] = overall[0]["sort_order"]
    elif mutation == "overall_sort_mismatch":
        overall[1]["sort_order"] = 9
    elif mutation == "duplicate_threshold":
        overall[1]["range_start_value"] = overall[0]["range_start_value"]
        levels[1]["score_band"] = "100%+"
    elif mutation == "out_of_range_threshold":
        overall[0]["range_start_value"] = 101
        levels[0]["score_band"] = "101%+"
    elif mutation == "threshold_multiplier_mismatch":
        overall[1]["range_start_value"] = 50
        levels[1]["score_band"] = "50%+"
    elif mutation == "score_band_mismatch":
        levels[1]["score_band"] = "59%+"

    source = tmp_path / f"{mutation}.json"
    source.write_text(json.dumps(data), encoding="utf-8")
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(source)
    assert {
        "RUBRICS_ADAPTER_UNTRANSLATABLE",
        "SOURCE_SCHEMA_UNSUPPORTED",
    } & diagnostic_codes(caught.value)
    serialized = json.dumps(caught.value.diagnostics)
    assert "private diagnostic" not in serialized
    assert "private rubric prose" not in serialized
    assert "private overall prose" not in serialized
    assert "private feedback" not in serialized


@pytest.mark.parametrize(
    "mutation",
    [
        "extra_activity_resource",
        "missing_resource",
        "duplicate_identifier",
        "external_href",
        "wrong_type",
        "wrong_material_type",
        "resource_dependency",
        "organizations",
    ],
)
def test_manifest_semantics_are_exact_for_folder_and_zip(
    tmp_path: Path,
    mutation: str,
) -> None:
    result = build_weave_outputs(
        FIXTURES / "three_level_explicit.md",
        tmp_path / f"built-{mutation}",
    )
    manifest_path = result["package_dir"] / "imsmanifest.xml"
    tree = ET.parse(manifest_path)
    root = tree.getroot()
    ims = "http://www.imsglobal.org/xsd/imscp_v1p1"
    d2l = "http://desire2learn.com/xsd/d2lcp_v2p0"
    resources = root.find(f"{{{ims}}}resources")
    assert resources is not None
    resource_nodes = resources.findall(f"{{{ims}}}resource")
    orgunit, rubrics = resource_nodes
    if mutation == "extra_activity_resource":
        ET.SubElement(
            resources,
            f"{{{ims}}}resource",
            {
                "identifier": "private_activity",
                "type": "webcontent",
                f"{{{d2l}}}material_type": "dropbox",
                f"{{{d2l}}}link_target": "",
                "href": "https://private.invalid/activity",
                "title": "",
            },
        )
    elif mutation == "missing_resource":
        resources.remove(rubrics)
    elif mutation == "duplicate_identifier":
        rubrics.set("identifier", orgunit.attrib["identifier"])
    elif mutation == "external_href":
        rubrics.set("href", "https://private.invalid/rubric")
    elif mutation == "wrong_type":
        rubrics.set("type", "imsqti_xmlv1p2")
    elif mutation == "wrong_material_type":
        rubrics.set(f"{{{d2l}}}material_type", "dropbox")
    elif mutation == "resource_dependency":
        ET.SubElement(rubrics, f"{{{ims}}}dependency", {"identifierref": "private"})
    else:
        ET.SubElement(root, f"{{{ims}}}organizations")
    tree.write(manifest_path, encoding="utf-8", xml_declaration=True)

    folder_errors, _, _ = validate_package_path(result["package_dir"])
    assert folder_errors
    mutated_zip = tmp_path / f"{mutation}.zip"
    rubric_authoring.zip_package(result["package_dir"], mutated_zip)
    zip_errors, _, _ = validate_package_path(mutated_zip)
    assert zip_errors
    assert "private.invalid" not in json.dumps(folder_errors + zip_errors)


@pytest.mark.parametrize(
    "legacy_hint",
    [
        "score_source",
        "weight_source",
        "grade_item_id",
        "kind",
        "source_title",
        "attachment",
        "association",
        "activity",
    ],
)
def test_legacy_json_refuses_provenance_and_activity_hints(
    tmp_path: Path,
    legacy_hint: str,
) -> None:
    data = _strict_builder_contract(2)
    rubric = data["rubrics"][0]
    if legacy_hint == "score_source":
        rubric["levels"][0]["score_source"] = "extracted_cell_points"
    elif legacy_hint == "weight_source":
        rubric["criteria"][0]["weight_source"] = "extracted_cell_points"
    else:
        rubric[legacy_hint] = "private hint"
    source = tmp_path / f"{legacy_hint}.json"
    source.write_text(json.dumps(data), encoding="utf-8")
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(source)
    assert "SOURCE_SCHEMA_UNSUPPORTED" in diagnostic_codes(caught.value)
    assert "private hint" not in json.dumps(caught.value.diagnostics)


def test_legacy_numeric_values_receive_only_explicit_provenance(tmp_path: Path) -> None:
    data = _strict_builder_contract(2)
    source = tmp_path / "explicit.json"
    source.write_text(json.dumps(data), encoding="utf-8")
    contract = normalize_source(source)
    rubric = contract["rubrics"][0]
    assert {item["score_source"] for item in rubric["levels"]} == {
        "explicit_level_metadata"
    }
    assert {item["weight_source"] for item in rubric["criteria"]} == {
        "explicit_weight"
    }


@pytest.mark.parametrize(
    "mutation",
    [
        "mixed_level_source",
        "mixed_weight_source",
        "mixed_overall_source",
        "self_asserted_extracted",
        "self_asserted_extracted_overall",
    ],
)
def test_declared_authoring_contract_refuses_mixed_or_untrusted_provenance(
    tmp_path: Path,
    mutation: str,
) -> None:
    contract = normalize_source(FIXTURES / "three_level_explicit.md")
    rubric = contract["rubrics"][0]
    if mutation == "mixed_level_source":
        rubric["levels"][0]["score_source"] = "approved_even_spacing"
    elif mutation == "mixed_weight_source":
        rubric["criteria"][0]["weight_source"] = "approved_equal_weights"
    elif mutation == "mixed_overall_source":
        rubric["overall_levels"][0]["source"] = "explicit_overall_threshold"
    elif mutation == "self_asserted_extracted":
        for level in rubric["levels"]:
            level["score_source"] = "extracted_cell_points"
        for criterion in rubric["criteria"]:
            criterion["weight_source"] = "extracted_cell_points"
    else:
        for overall in rubric["overall_levels"]:
            overall["source"] = "extracted_overall_threshold"
    source = tmp_path / f"{mutation}.json"
    source.write_text(json.dumps(contract), encoding="utf-8")
    with pytest.raises(AuthoringRefusal) as caught:
        normalize_source(
            source,
            allow_even_spacing=True,
            allow_equal_weights=True,
        )
    assert {
        "SCORING_METADATA_REQUIRED",
        "CRITERION_WEIGHT_REQUIRED",
        "RUBRICS_ADAPTER_UNTRANSLATABLE",
    } & diagnostic_codes(caught.value)

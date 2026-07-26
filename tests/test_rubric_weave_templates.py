from __future__ import annotations

import hashlib
import json
import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


REPO_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = REPO_ROOT / "scripts"
TEMPLATE_DIR = (
    REPO_ROOT / "workspace/reference/templates/rubric-weave/v1"
)
sys.path.insert(0, str(SCRIPTS))

import generate_rubric_weave_intake_templates as templates  # noqa: E402
from rubric_authoring import build_weave_outputs, normalize_source  # noqa: E402
from rubric_package_lib import validate_package_path  # noqa: E402


def _manifest() -> dict[str, object]:
    return json.loads(
        (TEMPLATE_DIR / templates.MANIFEST_NAME).read_text(encoding="utf-8")
    )


def _rubric_contract(path: Path) -> dict[str, object]:
    return normalize_source(path)["rubrics"][0]


def test_committed_assets_are_exact_deterministic_regeneration(
    tmp_path: Path,
) -> None:
    first = templates.write_assets(tmp_path / "first")
    second = templates.write_assets(tmp_path / "second")
    assert templates.check_assets(TEMPLATE_DIR) == []
    for name in (
        templates.DOCX_NAME,
        templates.MARKDOWN_NAME,
        templates.MANIFEST_NAME,
    ):
        canonical = (TEMPLATE_DIR / name).read_bytes()
        assert first[name].read_bytes() == canonical
        assert second[name].read_bytes() == canonical


def test_manifest_records_exact_version_media_path_bytes_and_sha256() -> None:
    manifest = _manifest()
    assert manifest["schema"] == "coursecraft.rubric_weave_template_manifest/1"
    assert manifest["template_set"] == "rubric-weave-intake"
    assert manifest["version"] == "v1"
    assert manifest["path_base"] == "manifest_directory"
    assert manifest["accepted_producer"] == {
        "repository": "coursecraft_workbench",
        "commit": "7c5140545548c89a254ac4502cfdd7ee6fb44255",
        "authoring_contract": "coursecraft.rubric_authoring/1",
    }

    entries = {entry["path"]: entry for entry in manifest["templates"]}
    assert set(entries) == {
        "rubric-weave-intake-template.docx",
        "rubric-weave-intake-template.md",
    }
    assert entries[templates.DOCX_NAME]["media_type"] == templates.DOCX_MEDIA_TYPE
    assert (
        entries[templates.MARKDOWN_NAME]["media_type"]
        == templates.MARKDOWN_MEDIA_TYPE
    )
    for name, entry in entries.items():
        path = TEMPLATE_DIR / name
        data = path.read_bytes()
        assert entry["version"] == "v1"
        assert entry["bytes"] == len(data)
        assert entry["sha256"] == hashlib.sha256(data).hexdigest()


def test_docx_uses_the_accepted_table_first_shape_and_explicit_geometry() -> None:
    path = TEMPLATE_DIR / templates.DOCX_NAME
    document = Document(path)
    body_kinds = [
        child.tag.rsplit("}", 1)[-1]
        for child in document.element.body.iterchildren()
        if child.tag != qn("w:sectPr")
    ]
    assert body_kinds == ["p", "tbl", "p", "p", "p", "p", "p"]
    assert len(document.tables) == 1
    assert len([p for p in document.paragraphs if p.style.name == "Heading 1"]) == 1
    assert document.paragraphs[0].text == templates.DEFAULT_SPEC.title
    assert all(
        paragraph.style.name == "Normal"
        and paragraph._p.pPr.numPr is None
        for paragraph in document.paragraphs[1:]
    )

    table = document.tables[0]
    expected_headers = [
        "Criterion",
        "Weight",
        "Ready to Share (100)",
        "Needs Revision (60)",
        "Not Yet Demonstrated (0)",
    ]
    assert [cell.text for cell in table.rows[0].cells] == expected_headers
    assert not table._tbl.xpath(".//w:tcPr/w:gridSpan")
    assert not table._tbl.xpath(".//w:tcPr/w:vMerge")
    assert not table._tbl.xpath(".//w:tbl/w:tr/w:tc/w:tbl")

    tbl_pr = table._tbl.tblPr
    assert tbl_pr.find(qn("w:tblW")).get(qn("w:w")) == "9360"
    assert tbl_pr.find(qn("w:tblInd")).get(qn("w:w")) == "120"
    grid_widths = [
        int(column.get(qn("w:w")))
        for column in table._tbl.tblGrid.findall(qn("w:gridCol"))
    ]
    assert grid_widths == [1800, 900, 2220, 2220, 2220]
    for row in table.rows:
        cell_widths = [
            int(cell._tc.tcPr.find(qn("w:tcW")).get(qn("w:w")))
            for cell in row.cells
        ]
        assert cell_widths == grid_widths

    instruction_text = "\n".join(
        paragraph.text for paragraph in document.paragraphs[1:]
    )
    assert "never invents scoring silently" in instruction_text
    assert "not a Brightspace import" in instruction_text
    assert "Activity attachment is a separate manual step" in instruction_text
    assert "Do not merge cells" in instruction_text
    assert "multiple header bands" in instruction_text
    assert "one criterion per row" in instruction_text


def test_docx_and_markdown_express_the_same_rubric_semantics() -> None:
    docx = _rubric_contract(TEMPLATE_DIR / templates.DOCX_NAME)
    markdown = _rubric_contract(TEMPLATE_DIR / templates.MARKDOWN_NAME)
    assert docx == markdown
    assert [level["name"] for level in docx["levels"]] == [
        "Ready to Share",
        "Needs Revision",
        "Not Yet Demonstrated",
    ]
    assert [level["multiplier"] for level in docx["levels"]] == [1.0, 0.6, 0.0]
    assert {level["score_source"] for level in docx["levels"]} == {
        "numeric_level_header"
    }
    assert [criterion["weight"] for criterion in docx["criteria"]] == [
        40.0,
        35.0,
        25.0,
    ]
    assert {criterion["weight_source"] for criterion in docx["criteria"]} == {
        "explicit_weight"
    }


def test_both_templates_preflight_and_build_valid_packages_without_fallback(
    tmp_path: Path,
) -> None:
    for name in (templates.DOCX_NAME, templates.MARKDOWN_NAME):
        source = TEMPLATE_DIR / name
        contract = normalize_source(source)
        assert contract["approvals"] == {
            "even_spacing": False,
            "equal_weights": False,
        }
        assert not any(
            diagnostic["code"]
            in {"EVEN_SPACING_APPROVED", "EQUAL_WEIGHTS_APPROVED"}
            for diagnostic in contract["diagnostics"]
        )

        result = build_weave_outputs(
            source,
            tmp_path / f"{source.stem}-{source.suffix.lstrip('.')}",
        )
        for package_path in (result["package_dir"], result["zip_path"]):
            errors, _, summary = validate_package_path(package_path)
            assert errors == []
            assert summary["rubric_count"] == 1
        with ZipFile(result["zip_path"]) as archive:
            assert set(archive.namelist()) == {
                "imsmanifest.xml",
                "orgunitconfig/orgunitconfig.xml",
                "rubrics_d2l.xml",
            }
        receipt = json.loads(result["receipt_path"].read_text(encoding="utf-8"))
        assert receipt["extensions"]["approvals"] == {
            "even_spacing": False,
            "equal_weights": False,
        }
        mapping = result["mapping_path"].read_text(encoding="utf-8")
        assert "Package scope: rubric objects only" in mapping
        assert "no activity attachment payloads" in mapping


def test_generator_supports_flexible_row_and_level_counts(
    tmp_path: Path,
) -> None:
    two_level = templates.TemplateSpec(
        title="SYNTHETIC TWO-LEVEL RUBRIC",
        levels=(
            templates.Level("Complete", 100),
            templates.Level("Revise", 0),
        ),
        criteria=(
            templates.Criterion(
                "Synthetic criterion",
                100,
                ("Synthetic complete description.", "Synthetic revise description."),
            ),
        ),
    )
    five_level = templates.TemplateSpec(
        title="SYNTHETIC FIVE-LEVEL RUBRIC",
        levels=(
            templates.Level("Exceptional", 100),
            templates.Level("Strong", 80),
            templates.Level("Developing", 55),
            templates.Level("Early", 25),
            templates.Level("Absent", 0),
        ),
        criteria=tuple(
            templates.Criterion(
                f"Synthetic criterion {index}",
                weight,
                tuple(
                    f"Synthetic row {index}, level {level_index}."
                    for level_index in range(1, 6)
                ),
            )
            for index, weight in enumerate((10, 20, 30, 40), start=1)
        ),
    )

    for label, spec, expected_levels, expected_rows in (
        ("two", two_level, 2, 1),
        ("five", five_level, 5, 4),
    ):
        written = templates.write_assets(tmp_path / label, spec)
        for name in (templates.DOCX_NAME, templates.MARKDOWN_NAME):
            contract = normalize_source(written[name])
            rubric = contract["rubrics"][0]
            assert len(rubric["levels"]) == expected_levels
            assert len(rubric["criteria"]) == expected_rows
            assert sum(row["weight"] for row in rubric["criteria"]) == 100
            assert contract["approvals"] == {
                "even_spacing": False,
                "equal_weights": False,
            }
